from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse, HttpResponse
from django.contrib.auth import login, logout, authenticate
from django.utils.http import url_has_allowed_host_and_scheme
from django.contrib.auth.decorators import login_required
from opencirt.models import Incident, Note, User, Message, GenericIoc, UserRole, Task, Action, Impact, SharedFile, AuditLog, PlatformSettings, CtiProvider, IncidentCategory, Campaign
from .utils import verify_permissions, user_is_incident_responder_orpublic, user_is_incident_responder, update_first_actions, sync_incident_times, get_incidents_by_day_and_severity
from .threat_intel import schedule_lookup, ELIGIBLE_TYPES as THREAT_INTEL_ELIGIBLE_TYPES
import json
import os
import re
import random
from datetime import timedelta, datetime
from django.utils import timezone
from django.template.loader import render_to_string
from xhtml2pdf import pisa
from io import BytesIO
from collections import Counter
from .report_generators import (
    parse_sections, parse_tlp, TLP_STYLES,
    DEFAULT_SECTIONS, ALL_SECTIONS,
    generate_markdown, generate_deep_json,
)
import csv
import io
from PIL import Image
from io import StringIO
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from django.core.cache import cache

_LOGIN_RATE_LIMIT = 10     # max failed attempts
_LOGIN_RATE_WINDOW = 900   # 15 minutes in seconds
_AI_REPHRASE_DAILY_LIMIT = 20

# ── Access control helpers ────────────────────────────────────────────────────

def _forbidden(request, reason='You do not have permission to perform this action.', incident=None):
    return render(request, '403.html', {'reason': reason, 'incident': incident}, status=403)

def custom_403(request, exception=None):
    reason = str(exception) if exception else 'You do not have permission to perform this action.'
    return render(request, '403.html', {'reason': reason}, status=403)

def _not_found(request, reason='The page or resource you requested does not exist.'):
    return render(request, '404.html', {'reason': reason}, status=404)

def custom_404(request, exception=None):
    reason = str(exception) if exception else 'The page or resource you requested does not exist.'
    return render(request, '404.html', {'reason': reason}, status=404)

# ── Audit log helpers ─────────────────────────────────────────────────────────

def _get_client_ip(request):
    xff = request.META.get('HTTP_X_FORWARDED_FOR')
    return xff.split(',')[0].strip() if xff else request.META.get('REMOTE_ADDR', '')

def _audit(incident, request, action, object_type, description):
    """Create an AuditLog entry. Never raises — logs silently on error."""
    try:
        AuditLog.objects.create(
            incident=incident,
            user=request.user if request.user.is_authenticated else None,
            ip_address=_get_client_ip(request),
            action=action,
            object_type=object_type,
            description=description,
        )
    except Exception:
        pass

# TLP → python-docx RGBColor mapping (used by download_incident_word)
TLP_COLORS = {
    'WHITE': RGBColor(80, 80, 80),
    'GREEN': RGBColor(40, 167, 69),
    'AMBER': RGBColor(253, 126, 20),
    'RED':   RGBColor(220, 53, 69),
}

@login_required(login_url='login')
def home(request):
    incidents_list = list(_accessible_incidents(request.user).order_by('-created_at'))

    total_time_to_detect, total_time_to_respond, total_duration, total_genericiocs, incident_count = 0, 0, 0, 0, 0

    for incident in incidents_list:
        if incident.time_to_detect:
            total_time_to_detect += incident.time_to_detect.total_seconds()
        if incident.time_to_respond:
            total_time_to_respond += incident.time_to_respond.total_seconds()
        if incident.duration:
            total_duration += incident.duration.total_seconds()
        total_genericiocs += incident.genericiocs.count()
        incident_count += 1

    kpis = [
        {'label': "Time to detect (TTD)",  'value': timedelta(seconds=total_time_to_detect / incident_count) if incident_count else timedelta()},
        {'label': "Time to respond (TTR)",  'value': timedelta(seconds=total_time_to_respond / incident_count) if incident_count else timedelta()},
        {'label': "Duration",               'value': timedelta(seconds=total_duration / incident_count) if incident_count else timedelta()},
        {'label': 'Iocs found',             'value': total_genericiocs},
    ]


    
    # Prepare graphs data
    # Pie chart
    piechart_data = {
        'labels': list(Counter(inc.status for inc in incidents_list).keys()),
        'values': list(Counter(inc.status for inc in incidents_list).values()),
    }
    # Get sorted unique dates
    timechart_data = get_incidents_by_day_and_severity()

    # Extract sorted dates
    dates = sorted(timechart_data.keys())

    # Extract unique severities
    severities = set()
    for sev_data in timechart_data.values():
        severities.update(sev_data.keys())
    severities = sorted(severities)  # Sort to maintain order

    # Prepare data for Chart.js
    severity_data = {sev: [timechart_data[date].get(sev, 0) for date in dates] for sev in severities}

    # Prepare dataset for Chart.js
    datasets = []
    for severity in severities:
        datasets.append({
            "label": severity,
            "data": severity_data[severity],
            # "borderColor": severity_colors.get(severity, "gray"),
        })

    context = {
        "labels": dates,
        "datasets": datasets
    }

    # Stats bar
    total_count = len(incidents_list)
    active_count = sum(1 for inc in incidents_list if inc.status in ('OPEN', 'IN_PROGRESS'))
    critical_count = sum(1 for inc in incidents_list if inc.severity == 'CRITICAL')
    total_iocs = sum(inc.genericiocs.count() for inc in incidents_list)

    # Severity distribution donut
    severity_counter = Counter(inc.severity for inc in incidents_list)
    severity_chart_data = {
        'labels': list(severity_counter.keys()),
        'values': list(severity_counter.values()),
    }

    from .choices_processor import INCIDENT_RESOLUTION_CHOICES

    return render(request, 'home.html', {
        'incidents': incidents_list,
        'user': request.user,
        'status_counts': piechart_data,
        'timechart_data': context,
        'kpis': kpis,
        'total_count': total_count,
        'active_count': active_count,
        'critical_count': critical_count,
        'total_iocs': total_iocs,
        'severity_chart_data': severity_chart_data,
        'INCIDENT_RESOLUTION_CHOICES': INCIDENT_RESOLUTION_CHOICES,
        'INCIDENT_STATUS_CHOICES': [
            ('OPEN', 'Open'), ('IN_PROGRESS', 'In Progress'),
            ('RESOLVED', 'Resolved'), ('CLOSED', 'Closed'),
        ],
        'INCIDENT_SEVERITY_CHOICES': [
            ('LOW', 'Low'), ('MEDIUM', 'Medium'),
            ('HIGH', 'High'), ('CRITICAL', 'Critical'),
        ],
        'campaigns': Campaign.objects.all().order_by('name'),
        'categories': IncidentCategory.objects.all().order_by('name'),
    })


def _login_rate_limit_exceeded(ip: str) -> bool:
    """Return True if this IP has exceeded the failed-login rate limit."""
    key = f'login_fail_{ip}'
    try:
        count = cache.incr(key)
    except ValueError:
        cache.set(key, 1, timeout=_LOGIN_RATE_WINDOW)
        count = 1
    return count > _LOGIN_RATE_LIMIT


def _login_rate_limit_reset(ip: str) -> None:
    """Clear the failed-login counter for an IP on successful login."""
    cache.delete(f'login_fail_{ip}')


def custom_login(request):
    if request.method == 'POST':
        ip = _get_client_ip(request)
        if _login_rate_limit_exceeded(ip):
            return render(request, 'login.html',
                          {'error': 'Too many failed attempts. Try again in 15 minutes.'},
                          status=429)
        username = request.POST.get('username', '').strip()
        password = request.POST.get('password', '')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            _login_rate_limit_reset(ip)
            login(request, user)
            next_url = request.GET.get('next', '/home')
            if not url_has_allowed_host_and_scheme(
                next_url,
                allowed_hosts={request.get_host()},
                require_https=request.is_secure(),
            ):
                next_url = '/home'
            return redirect(next_url)
        return render(request, 'login.html', {'error': 'Invalid username or password'})
    return render(request, 'login.html')

def custom_logout(request):
    logout(request)
    return redirect('login')

@login_required(login_url='login')
def about(request):
    return render(request,'about.html', {'user': request.user})


@login_required(login_url='login')
def profile(request):
    try:
        current_user = request.user
    except:
        return JsonResponse({'error': 'La galere'}, status=404)
 
    prefs = request.user.preferences if request.user.preferences else {}
    return render(request, 'profile.html', {
        'user': request.user,
        'prefs': prefs,
    })



@login_required(login_url='login')
@user_is_incident_responder_orpublic
def overview(request, id):
    incident = get_object_or_404(Incident, id=id)
    user_role = _get_user_role(request.user, incident)
    if user_role is None:
        return _forbidden(request, 'You are not a member of this incident.', incident=incident)

    platform = PlatformSettings.get()
    return render(request, 'incidents/overview.html', {
        'incident':           incident,
        'user':               request.user,
        'current_user_role':  user_role,
        'ai_configured':      platform.ai_provider != 'NONE' and bool(platform.ai_api_key),
        'all_categories':     list(IncidentCategory.objects.order_by('name').values('id', 'name', 'color')),
    })

@login_required(login_url='login')
@user_is_incident_responder_orpublic
def activity(request, id):
    incident = get_object_or_404(Incident, pk=id)
    user_role = _get_user_role(request.user, incident)
    if user_role is None:
        return _forbidden(request, 'You are not a member of this incident.', incident=incident)
    incident_leads = User.objects.filter(id__in=UserRole.objects.filter(incident=incident, role="INCIDENT_LEAD").values_list('user_id', flat=True))
    return render(request,'incidents/activity.html', {'incident': incident, 'user': request.user,'current_user_role': user_role, 'incident_leads': incident_leads})

@login_required(login_url='login')
@user_is_incident_responder_orpublic
def impacts(request, id):
    incident = get_object_or_404(Incident, pk=id)
    user_role = _get_user_role(request.user, incident)
    if user_role is None:
        return _forbidden(request, 'You are not a member of this incident.', incident=incident)
    incident_leads = User.objects.filter(id__in=UserRole.objects.filter(incident=incident, role="INCIDENT_LEAD").values_list('user_id', flat=True))
    return render(request,'incidents/impacts.html', {'incident': incident, 'user': request.user,'current_user_role': user_role, 'incident_leads': incident_leads})

@login_required(login_url='login')
@user_is_incident_responder_orpublic
def notes(request, id):
    incident = get_object_or_404(Incident, pk=id)
    user_role = _get_user_role(request.user, incident)
    if user_role is None:
        return _forbidden(request, 'You are not a member of this incident.', incident=incident)
    shared_files = list(incident.shared_files.all().order_by('-created_at'))
    for sf in shared_files:
        sf.is_executable = _get_extension(sf.original_name) in _EXECUTABLE_EXTENSIONS
    return render(request, 'incidents/notes.html', {
        'incident':     incident,
        'user':         request.user,
        'current_user_role': user_role,
        'shared_files': shared_files,
    })

@login_required(login_url='login')
@user_is_incident_responder_orpublic
def tasks(request, id):
    incident = get_object_or_404(Incident, pk=id)
    user_role = _get_user_role(request.user, incident)
    if user_role is None:
        return _forbidden(request, 'You are not a member of this incident.', incident=incident)
    return render(request,'incidents/tasks.html', {'incident': incident, 'user': request.user, 'current_user_role': user_role})

@login_required(login_url='login')
@user_is_incident_responder
def add_note(request, id):

    try:
        if request.method == 'POST':
            incident = Incident.objects.get(pk=id)
            data = json.loads(request.body)
            name = data.get('name', 'My note')
            text = data.get('text', '')
            new_note = Note.objects.create(
            incident=incident,
            name=name,
            text=text,
            created_by = request.user
            )
            Message.objects.create(
                incident=incident,
                sender=request.user,
                text=f"{request.user.username} created note: {new_note.name}",
                is_bot=True,
                link=f"/incident/{incident.id}/notes"
            )
            _audit(incident, request, 'CREATE', 'Note', f'Created note "{new_note.name}"')
            return JsonResponse({
            'id': new_note.id,
            'name': new_note.name,
            'text': new_note.text
            })

    except Incident.DoesNotExist:
        return JsonResponse({'error': 'Invalid request'}, status=400)

@login_required(login_url='login')
@user_is_incident_responder
def update_note(request, id):
    if request.method == 'POST':
        try:
            
            data = json.loads(request.body)
            note_id = data.get('note_id')
            note = Note.objects.get(id=note_id)
            note.name = data.get('name', 'My note')
            note.text = data.get('text', '')
            
            note.save()
            _audit(Incident.objects.get(pk=id), request, 'UPDATE', 'Note', f'Updated note "{note.name}"')
            return JsonResponse({'status': 'success', 'note_id': note_id})

        except Note.DoesNotExist:
            return JsonResponse({'error': 'Note not found'}, status=404)
    return JsonResponse({'error': 'Invalid request'}, status=400)

@login_required(login_url='login')
@user_is_incident_responder
def delete_note(request, id):
    if request.method == 'DELETE':
        try:
            data = json.loads(request.body)
            note_id = data.get('note_id')
            note = Note.objects.get(pk=note_id)
            note_name = note.name
            note.delete()
            _audit(Incident.objects.get(pk=id), request, 'DELETE', 'Note', f'Deleted note "{note_name}"')
            return JsonResponse({"status": "success", "message": "Note deleted successfully", "note_id": note_id })
        except Note.DoesNotExist:
            return JsonResponse({'error': 'Note not found'}, status=404)        

# ── File upload security ──────────────────────────────────────────────────────
# Extensions that can be executed server-side (PHP, ASP, CGI, shell scripts…).
# These are NEVER accepted — an attacker could trigger server-side execution if
# the web server ever serves them directly.
_BLOCKED_EXTENSIONS = frozenset([
    # Web-side scripts
    'php', 'php3', 'php4', 'php5', 'php7', 'php8', 'phtml', 'phar',
    'asp', 'aspx', 'asa', 'asax', 'ascx', 'ashx', 'asmx',
    'jsp', 'jspx', 'jws', 'cgi',
    # Unix/Linux shells (could run via CGI or misconfigured handler)
    'sh', 'bash', 'zsh', 'fish', 'ksh', 'csh', 'tcsh',
    # Interpreted languages runnable server-side
    'py', 'pyc', 'pyo', 'rb', 'pl', 'lua', 'tcl',
    # Web-server config overrides
    'htaccess', 'htpasswd',
])

# Client-side executables / binaries — safe for the server to store, but may
# be malicious in content.  Uploads are accepted; downloads require confirmation.
_EXECUTABLE_EXTENSIONS = frozenset([
    # Windows executables & installers
    'exe', 'com', 'scr', 'pif', 'msi', 'msp', 'msc', 'cpl',
    # Windows scripting
    'bat', 'cmd', 'vbs', 'vbe', 'js', 'jse', 'ws', 'wsf', 'wsc', 'wsh',
    'ps1', 'psm1', 'psd1', 'ps1xml', 'lnk', 'url', 'hta',
    # Compiled binaries / libraries
    'dll', 'so', 'dylib', 'elf',
    # JVM bytecode / packages
    'jar', 'war', 'ear', 'class',
    # Mobile packages
    'apk', 'ipa', 'xap',
])

_MAX_UPLOAD_BYTES = 50 * 1024 * 1024  # 50 MB


def _sanitize_filename(name):
    """
    Return a safe filename:
    - Strip any path components (defence against path traversal)
    - Remove null bytes and control characters
    - Truncate to 255 characters
    - Fall back to 'file' if nothing is left
    """
    name = os.path.basename(name.replace('\\', '/'))          # strip path
    name = re.sub(r'[\x00-\x1f\x7f/<>:"|?*]', '_', name)     # remove control chars & OS-reserved chars
    name = name.strip('. ')                                    # no leading/trailing dots or spaces
    return name[:255] or 'file'


def _get_extension(name):
    """Return lowercase extension without leading dot, e.g. 'pdf'."""
    return os.path.splitext(name)[1].lstrip('.').lower()


# ─────────────────────────────────────────────────────────────────────────────

@login_required(login_url='login')
@user_is_incident_responder
def upload_file(request, id):
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)

    incident = get_object_or_404(Incident, pk=id)
    f = request.FILES.get('file')
    if not f:
        return JsonResponse({'error': 'No file provided'}, status=400)

    # Size check
    if f.size > _MAX_UPLOAD_BYTES:
        return JsonResponse({'error': 'File too large (max 50 MB)'}, status=400)

    # Sanitize the filename — never trust user input for paths
    safe_name = _sanitize_filename(f.name)
    ext = _get_extension(safe_name)

    # Block server-side-executable scripts (PHP, ASP, shell, etc.)
    if ext in _BLOCKED_EXTENSIONS:
        return JsonResponse(
            {'error': f'File type ".{ext}" cannot be uploaded — server-side scripts are not permitted.'},
            status=400
        )

    # Overwrite the in-memory name so Django's FileField stores it safely
    f.name = safe_name

    shared = SharedFile.objects.create(
        incident=incident,
        uploaded_by=request.user,
        file=f,
        original_name=safe_name,
        size=f.size,
    )
    _audit(incident, request, 'CREATE', 'File', f'Uploaded file "{safe_name}" ({f.size} bytes)')
    return JsonResponse({
        'status': 'success',
        'id': shared.id,
        'name': shared.original_name,
        'size': shared.size,
        'url': shared.file.url,
        'uploaded_by': request.user.username,
        'created_at': shared.created_at.strftime('%d %b %Y, %H:%M'),
    })


@login_required(login_url='login')
@user_is_incident_responder
def delete_file(request, id):
    if request.method != 'DELETE':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        data = json.loads(request.body)
        file_id = data.get('file_id')
        incident = get_object_or_404(Incident, pk=id)
        shared = get_object_or_404(SharedFile, pk=file_id, incident=incident)
        file_name = shared.original_name
        shared.file.delete(save=False)
        shared.delete()
        _audit(incident, request, 'DELETE', 'File', f'Deleted file "{file_name}"')
        return JsonResponse({'status': 'success'})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@login_required(login_url='login')
@user_is_incident_responder_orpublic
def incident_settings(request, id):
    incident = get_object_or_404(Incident, pk=id)
    user_role = _get_user_role(request.user, incident)
    if user_role is None:
        return _forbidden(request, 'You are not a member of this incident.', incident=incident)
    pr = getattr(request.user, 'platform_role', '') or ''
    if user_role.role not in ('INCIDENT_LEAD',) and not request.user.is_admin and pr != 'SOC_LEAD':
        return _forbidden(request, 'Only Incident Leads and platform admins can access incident settings.', incident=incident)

    platform = PlatformSettings.get()
    all_incidents = _accessible_incidents(request.user).order_by('name')
    from .choices_processor import INCIDENT_RESOLUTION_CHOICES
    return render(request, 'incidents/incident_settings.html', {
        'incident':                    incident,
        'user':                        request.user,
        'user_role':                   user_role.role,
        'current_user_role':           user_role,
        'ai_configured':               platform.ai_provider != 'NONE' and bool(platform.ai_api_key),
        'all_incidents':               all_incidents,
        'INCIDENT_RESOLUTION_CHOICES': INCIDENT_RESOLUTION_CHOICES,
    })

@login_required(login_url='login')
@user_is_incident_responder_orpublic
def report_preview(request, id):
    """
    GET /api/incident/<id>/report-preview/?sections=executive_summary,iocs&tlp=AMBER
    Returns HTML string for iframe srcdoc.
    """
    try:
        incident = Incident.objects.get(pk=id)
    except Incident.DoesNotExist:
        return HttpResponse('<p style="padding:20px;color:#dc2626;">Incident not found.</p>', status=404)

    sections = parse_sections(request.GET)
    tlp = parse_tlp(request.GET)
    tlp_style = TLP_STYLES[tlp]

    html = render_to_string('reports/report_template.html', {
        'incident': incident,
        'sections': sections,
        'tlp': tlp,
        'tlp_style': tlp_style,
        'is_pdf': False,
        'generated_at': timezone.now().strftime('%d %B %Y, %H:%M'),
        'generated_by': request.user.username,
    })
    return HttpResponse(html, content_type='text/html; charset=utf-8')


@login_required(login_url='login')
@user_is_incident_responder_orpublic
def report(request, id):
    incident = get_object_or_404(Incident, pk=id)
    user_role = _get_user_role(request.user, incident)
    if user_role is None:
        return _forbidden(request, 'You are not a member of this incident.', incident=incident)

    return render(request, 'incidents/report.html', {
        'incident': incident,
        'user': request.user,
        'current_user_role': user_role,
        'all_sections': list(ALL_SECTIONS),
        'default_sections': list(DEFAULT_SECTIONS),
    })

@login_required(login_url='login')
@user_is_incident_responder_orpublic
def iocs(request, id):
    incident = get_object_or_404(Incident, pk=id)
    user_role = _get_user_role(request.user, incident)
    if user_role is None:
        return _forbidden(request, 'You are not a member of this incident.', incident=incident)
    cti_configured      = CtiProvider.objects.filter(enabled=True).exclude(api_key='').exists()
    supported_ioc_types = _cti_supported_types()
    return render(request, 'incidents/evidence.html', {
        'incident':            incident,
        'user':                request.user,
        'current_user_role':   user_role,
        'cti_configured':      cti_configured,
        'supported_ioc_types': supported_ioc_types,
    })

@login_required(login_url='login')
@user_is_incident_responder_orpublic
def timeline(request, id):
    incident = get_object_or_404(Incident, pk=id)
    user_role = _get_user_role(request.user, incident)
    if user_role is None:
        return _forbidden(request, 'You are not a member of this incident.', incident=incident)
    return render(request,'incidents/timeline.html', {'incident': incident, 'user': request.user, 'current_user_role': user_role})

def join(request, id):
    try:
        incident = Incident.objects.get(pk=id)
    except Incident.DoesNotExist:
        return redirect('/home')

    if request.method == 'POST':
        code = request.POST.get('code', '').strip()

        # Ensure every incident has a code (backfill for any created before this guard)
        if not incident.invite_code:
            incident.invite_code = generate_invite_code()
            incident.save(update_fields=['invite_code'])

        if code != incident.invite_code:
            return render(request, 'incidents/join.html', {
                'incident': incident,
                'error': 'Invalid code. Please check with your incident lead.',
                'user': request.user,
            })

        if request.user.is_authenticated:
            # Already logged in — just add to incident
            if not UserRole.objects.filter(user=request.user, incident=incident).exists():
                UserRole.objects.create(
                    user=request.user,
                    incident=incident,
                    role='READER',
                    display_role='Responder'
                )
            return redirect('overview', id=incident.id)
        else:
            # Create a new account
            username = request.POST.get('username', '').strip()
            email = request.POST.get('email', '').strip()
            password = request.POST.get('password', '').strip()

            if not username or not password:
                return render(request, 'incidents/join.html', {
                    'incident': incident,
                    'error': 'Username and password are required.',
                })
            if User.objects.filter(username=username).exists():
                return render(request, 'incidents/join.html', {
                    'incident': incident,
                    'error': 'Username already taken. Please choose another.',
                })
            user = User.objects.create_user(username=username, email=email, password=password)
            UserRole.objects.create(
                user=user,
                incident=incident,
                role='READER',
                display_role='Responder'
            )
            login(request, user)
            return redirect('overview', id=incident.id)

    return render(request, 'incidents/join.html', {
        'incident': incident,
        'user': request.user,
    })


# ─────────────────────────────────────────────────────
# INVITE HELPERS
# ─────────────────────────────────────────────────────

def generate_invite_code():
    return str(random.randint(100000, 999999))


# ─────────────────────────────────────────────────────
# CREATE INCIDENT
# ─────────────────────────────────────────────────────

@login_required(login_url='login')
def create_incident(request):
    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        description = request.POST.get('description', '').strip()
        severity = request.POST.get('severity', 'MEDIUM')
        starting_time_str = request.POST.get('starting_time', '')
        is_public = request.POST.get('is_public') == 'on'

        if not name:
            return render(request, 'incidents/create_incident.html', {
                'user': request.user,
                'error': 'Incident name is required.',
                'form_data': request.POST,
            })

        # Parse start time
        if starting_time_str:
            try:
                starting_time = datetime.strptime(starting_time_str, '%Y-%m-%dT%H:%M')
                starting_time = timezone.make_aware(starting_time)
            except ValueError:
                starting_time = timezone.now()
        else:
            starting_time = timezone.now()

        invite_code = generate_invite_code()

        incident = Incident.objects.create(
            name=name,
            description=description,
            severity=severity,
            status='OPEN',
            starting_time=starting_time,
            ending_time=starting_time,
            duration=timedelta(0),
            time_to_detect=timedelta(0),
            time_to_respond=timedelta(0),
            executive_summary='',
            lessons_learned='',
            technical_details='',
            is_public=is_public,
            created_by=request.user,
            invite_code=invite_code,
        )

        UserRole.objects.create(
            user=request.user,
            incident=incident,
            role='INCIDENT_LEAD',
            display_role='Incident Lead',
        )

        return redirect('incident_invite', id=incident.id)

    return render(request, 'incidents/create_incident.html', {'user': request.user})


# ─────────────────────────────────────────────────────
# POST-CREATION INVITE SCREEN
# ─────────────────────────────────────────────────────

@login_required(login_url='login')
def incident_invite(request, id):
    try:
        incident = Incident.objects.get(pk=id)
    except Incident.DoesNotExist:
        return redirect('/home')

    if not incident.invite_code:
        incident.invite_code = generate_invite_code()
        incident.save()

    join_url = request.build_absolute_uri(f'/incident/{id}/join')

    return render(request, 'incidents/invite.html', {
        'incident': incident,
        'user': request.user,
        'join_url': join_url,
    })


@login_required(login_url='login')
def regenerate_invite(request, id):
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        incident = Incident.objects.get(pk=id)
        user_role = _get_user_role(request.user, incident)
        if user_role is None or user_role.role not in ('INCIDENT_LEAD', 'RESPONDER'):
            return JsonResponse({'error': 'Permission denied'}, status=403)
        incident.invite_code = generate_invite_code()
        incident.save()
        return JsonResponse({'status': 'success', 'code': incident.invite_code})
    except Incident.DoesNotExist:
        return JsonResponse({'error': 'Incident not found'}, status=404)


# ─────────────────────────────────────────────────────
# SETTINGS PAGE
# ─────────────────────────────────────────────────────

@login_required(login_url='login')
def settings_view(request):
    if not request.user.is_admin:
        return _forbidden(request, 'Platform settings are restricted to administrators.')

    user = request.user
    prefs = user.preferences if user.preferences else {}

    if request.method == 'POST':
        prefs['default_severity'] = request.POST.get('default_severity', 'MEDIUM')
        prefs['chart_period'] = request.POST.get('chart_period', '30d')
        user.preferences = prefs
        user.save()
        return redirect('/settings?saved=1')

    platform      = PlatformSettings.get()
    cti_providers = CtiProvider.objects.all().order_by('name')
    all_users     = User.objects.order_by('username')
    return render(request, 'settings.html', {
        'user':          user,
        'prefs':         prefs,
        'platform':      platform,
        'cti_providers': cti_providers,
        'all_users':     all_users,
    })


@login_required(login_url='login')
def api_admin_set_platform_role(request, user_id):
    """POST /api/admin/users/<user_id>/set-platform-role/ — admin only."""
    if not request.user.is_admin:
        return JsonResponse({'error': 'Admin only.'}, status=403)
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed.'}, status=405)
    try:
        target = User.objects.get(pk=user_id)
    except User.DoesNotExist:
        return JsonResponse({'error': 'User not found.'}, status=404)

    try:
        data = json.loads(request.body)
    except (ValueError, KeyError):
        return JsonResponse({'error': 'Invalid JSON.'}, status=400)

    new_role = data.get('platform_role', '')
    valid = {'', 'SOC_ANALYST', 'SOC_LEAD'}
    if new_role not in valid:
        return JsonResponse({'error': 'Invalid platform_role.'}, status=400)

    target.platform_role = new_role
    target.save(update_fields=['platform_role'])
    return JsonResponse({'status': 'ok', 'platform_role': new_role})


@login_required(login_url='login')
@user_is_incident_responder_orpublic
def download_incident_json(request, id):
    """
    GET /api/incident/<id>/download-json/
    Returns full deep JSON export (sections don't apply — always full).
    """
    try:
        incident = Incident.objects.get(pk=id)
    except Incident.DoesNotExist:
        return JsonResponse({'error': 'Incident not found'}, status=404)

    tlp = parse_tlp(request.GET)
    data = generate_deep_json(incident, generated_by=request.user.username, tlp=tlp)

    response = HttpResponse(
        json.dumps(data, indent=2, ensure_ascii=False),
        content_type='application/json; charset=utf-8'
    )
    response['Content-Disposition'] = (
        f'attachment; filename="incident_{incident.id}_full.json"'
    )
    return response


@login_required(login_url='login')
@user_is_incident_responder_orpublic
def download_incident_csv(request, id):
    """
    GET /api/incident/<id>/download-csv/
    Returns a CSV of all IoCs (sections don't apply).
    Columns: Type, Value, Status, Description, Created At, Linked Actions
    """
    try:
        incident = Incident.objects.get(pk=id)
    except Incident.DoesNotExist:
        return JsonResponse({'error': 'Incident not found'}, status=404)

    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(['Type', 'Value', 'Status', 'Threat Intel Verdict', 'Description', 'Created At', 'Linked Actions'])

    for ioc in (
        incident.genericiocs.all()
        .prefetch_related('actions')
        .select_related('created_by')
    ):
        linked = ', '.join(str(a.title) for a in ioc.actions.all())
        rep = ioc.reputation
        if rep:
            verdict = rep.get('status', 'unknown').upper()
            vt = rep.get('vt') or {}
            if vt.get('total'):
                verdict += f" ({vt.get('malicious', 0)}/{vt['total']} engines)"
        else:
            verdict = ''
        writer.writerow([
            ioc.get_type_display(),
            ioc.value,
            ioc.get_status_display(),
            verdict,
            ioc.description or '',
            ioc.created_at.strftime('%Y-%m-%d %H:%M') if ioc.created_at else '',
            linked,
        ])

    response = HttpResponse(output.getvalue(), content_type='text/csv; charset=utf-8')
    response['Content-Disposition'] = (
        f'attachment; filename="incident_{incident.id}_iocs.csv"'
    )
    return response


@login_required(login_url='login')
@user_is_incident_responder
def download_incident_html(request, id):
    """
    POST /api/incident/<id>/download-html/
    Body: sections=..., tlp=...
    Returns the report as a self-contained .html archive download.
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    try:
        incident = Incident.objects.get(pk=id)
    except Incident.DoesNotExist:
        return JsonResponse({'error': 'Incident not found'}, status=404)

    sections = parse_sections(request.POST)
    tlp = parse_tlp(request.POST)
    tlp_style = TLP_STYLES[tlp]

    html = render_to_string('reports/report_template.html', {
        'incident': incident,
        'sections': sections,
        'tlp': tlp,
        'tlp_style': tlp_style,
        'is_pdf': False,
        'generated_at': timezone.now().strftime('%d %B %Y, %H:%M'),
        'generated_by': request.user.username,
    })

    response = HttpResponse(html, content_type='text/html; charset=utf-8')
    response['Content-Disposition'] = (
        f'attachment; filename="incident_{incident.id}_report.html"'
    )
    return response

@login_required(login_url='login')
@user_is_incident_responder
def download_incident_markdown(request, id):
    """
    POST /api/incident/<id>/download-markdown/
    Body: sections=..., tlp=...
    Returns a .md file download.
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    try:
        incident = Incident.objects.get(pk=id)
    except Incident.DoesNotExist:
        return JsonResponse({'error': 'Incident not found'}, status=404)

    sections = parse_sections(request.POST)
    tlp = parse_tlp(request.POST)

    content = generate_markdown(incident, sections, tlp, request.user.username)

    response = HttpResponse(content, content_type='text/markdown; charset=utf-8')
    response['Content-Disposition'] = (
        f'attachment; filename="incident_{incident.id}_report.md"'
    )
    return response

@login_required(login_url='login')
@user_is_incident_responder
def download_incident_pdf(request, id):
    """
    POST /api/incident/<id>/download-pdf/
    Body: sections=..., tlp=...
    Returns a PDF file download using xhtml2pdf.
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    try:
        incident = Incident.objects.get(pk=id)
    except Incident.DoesNotExist:
        return JsonResponse({'error': 'Incident not found'}, status=404)

    sections = parse_sections(request.POST)
    tlp = parse_tlp(request.POST)
    tlp_style = TLP_STYLES[tlp]

    html = render_to_string('reports/report_template.html', {
        'incident': incident,
        'sections': sections,
        'tlp': tlp,
        'tlp_style': tlp_style,
        'is_pdf': True,
        'generated_at': timezone.now().strftime('%d %B %Y, %H:%M'),
        'generated_by': request.user.username,
    })

    buffer = BytesIO()
    pisa_status = pisa.CreatePDF(html, dest=buffer, encoding='utf-8')

    if pisa_status.err:
        return HttpResponse(
            f'PDF generation error (xhtml2pdf code {pisa_status.err}). '
            f'Try the HTML export as a workaround.',
            status=500
        )

    buffer.seek(0)
    response = HttpResponse(buffer.read(), content_type='application/pdf')
    response['Content-Disposition'] = (
        f'attachment; filename="incident_{incident.id}_report.pdf"'
    )
    return response

    
@login_required(login_url='login')
@user_is_incident_responder
def download_incident_word(request, id):
    """
    POST /api/incident/<id>/download-word/
    Body: sections=..., tlp=...
    Returns a .docx file download.
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    try:
        incident = Incident.objects.get(pk=id)
    except Incident.DoesNotExist:
        return JsonResponse({'error': 'Incident not found'}, status=404)

    sections = parse_sections(request.POST)
    tlp = parse_tlp(request.POST)

    doc = DocxDocument()

    # ── Cover ──
    doc.add_heading(incident.name, 0)

    tlp_para = doc.add_paragraph()
    tlp_run = tlp_para.add_run(f'TLP:{tlp}')
    tlp_run.bold = True
    tlp_run.font.size = Pt(13)
    tlp_run.font.color.rgb = TLP_COLORS.get(tlp, RGBColor(80, 80, 80))

    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(
        f'Generated {timezone.now().strftime("%d %B %Y, %H:%M")} by {request.user.username}'
    )
    meta_run.font.color.rgb = RGBColor(130, 130, 130)
    doc.add_paragraph()  # spacer

    # ── Executive Summary ──
    if 'executive_summary' in sections:
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(incident.executive_summary or 'No executive summary provided.')

    # ── Metadata ──
    if 'metadata' in sections:
        doc.add_heading('Incident Metadata', 1)
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        for label, value in [
            ('Severity', incident.severity),
            ('Status', incident.get_status_display()),
            ('Start time', str(incident.starting_time)),
            ('End time', str(incident.ending_time)),
            ('Duration', str(incident.duration)),
            ('Time to detect', str(incident.time_to_detect)),
            ('Time to respond', str(incident.time_to_respond)),
            ('Created by', incident.created_by.username if incident.created_by else 'Unknown'),
            ('Public', 'Yes' if incident.is_public else 'No'),
        ]:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value

    # ── Responders ──
    if 'responders' in sections:
        doc.add_heading('Responders', 1)
        roles_qs = incident.incident_roles.all().select_related('user')
        if roles_qs.exists():
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            for i, h in enumerate(['Username', 'Display Name', 'Role', 'Display Role']):
                table.rows[0].cells[i].text = h
            for ur in roles_qs:
                row = table.add_row().cells
                row[0].text = ur.user.username
                row[1].text = ur.user.displayname or '-'
                row[2].text = ur.get_role_display()
                row[3].text = ur.display_role or '-'
        else:
            doc.add_paragraph('No responders recorded.')

    # ── Timeline ──
    if 'timeline' in sections:
        doc.add_heading('Timeline', 1)
        actions_qs = incident.actions.all().order_by('observed_at').select_related('created_by')
        if not actions_qs.exists():
            doc.add_paragraph('No timeline events recorded.')
        else:
            for action in actions_qs:
                if action.observed_at:
                    time_str = action.observed_at.strftime('%d %b %Y %H:%M')
                elif action.starting_time:
                    time_str = action.starting_time.strftime('%d %b %Y %H:%M')
                else:
                    time_str = ''
                p = doc.add_paragraph(style='List Bullet')
                r = p.add_run(f'[{action.get_type_display()}] {action.title}')
                r.bold = True
                if time_str:
                    p.add_run(f'  —  {time_str}')
                if action.description:
                    doc.add_paragraph(action.description)

    # ── IoCs ──
    if 'iocs' in sections:
        doc.add_heading('IoC / Evidence', 1)
        iocs_qs = incident.genericiocs.all()
        if not iocs_qs.exists():
            doc.add_paragraph('No IoCs recorded.')
        else:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            for i, h in enumerate(['Type', 'Value', 'Status', 'Threat Intel Verdict', 'Description']):
                table.rows[0].cells[i].text = h
            for ioc in iocs_qs:
                rep = ioc.reputation
                if rep:
                    verdict = rep.get('status', 'unknown').upper()
                    vt = rep.get('vt') or {}
                    if vt.get('total'):
                        verdict += f" ({vt.get('malicious', 0)}/{vt['total']} engines)"
                else:
                    verdict = '—'
                row = table.add_row().cells
                row[0].text = ioc.get_type_display()
                row[1].text = ioc.value
                row[2].text = ioc.get_status_display()
                row[3].text = verdict
                row[4].text = ioc.description or '-'

    # ── Tasks ──
    if 'tasks' in sections:
        doc.add_heading('Tasks', 1)
        tasks_qs = incident.tasks.all().select_related('assignee')
        if not tasks_qs.exists():
            doc.add_paragraph('No tasks recorded.')
        else:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            for i, h in enumerate(['Priority', 'Title', 'Status', 'Assignee', 'Description']):
                table.rows[0].cells[i].text = h
            for task in tasks_qs:
                row = table.add_row().cells
                row[0].text = task.priority
                row[1].text = task.title
                row[2].text = task.status
                row[3].text = task.assignee.username if task.assignee else '-'
                row[4].text = task.description or '-'

    # ── Notes ──
    if 'notes' in sections:
        doc.add_heading('Notes', 1)
        notes_qs = incident.notes.all().select_related('created_by')
        if not notes_qs.exists():
            doc.add_paragraph('No notes recorded.')
        else:
            for note in notes_qs:
                doc.add_heading(note.name, 2)
                author = note.created_by.username if note.created_by else 'Unknown'
                p = doc.add_paragraph()
                p.add_run(
                    f'{author}  ·  {note.created_at.strftime("%d %b %Y %H:%M") if note.created_at else ""}'
                ).italic = True
                doc.add_paragraph(note.text)

    # ── Lessons Learned ──
    if 'lessons_learned' in sections:
        doc.add_heading('Lessons Learned', 1)
        ll = incident.lessons_learned
        doc.add_paragraph(
            ll if ll and ll != 'SOME STRING' else 'No lessons learned recorded.'
        )

    # ── Technical Details ──
    if 'technical_details' in sections:
        doc.add_heading('Technical Details', 1)
        td = incident.technical_details
        doc.add_paragraph(
            td if td and td != 'SOME STRING' else 'No technical details recorded.'
        )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = (
        f'attachment; filename="incident_{incident.id}_report.docx"'
    )
    return response


@login_required
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
@user_is_incident_responder
def send_message(request, id):
    if request.method == 'POST':
        data = json.loads(request.body)
        text = data.get('message')
        incident = Incident.objects.get(pk=id)
        
        # Create a new message in the database
        message = Message.objects.create(
            incident=incident,
            sender=request.user,
            text=text
        )

        return JsonResponse({
            'status': 'success',
            'message': message.text,
            'sender': message.sender.username,
            'timestamp': message.created_at
        })

    return JsonResponse({'status': 'error', 'message': 'Invalid request'}, status=400)


@login_required(login_url='login')
@user_is_incident_responder
def get_messages(request, id):
    try:
        incident = Incident.objects.get(pk=id)
        msgs = incident.messages.all().order_by('created_at').select_related('sender')
        data = []
        for msg in msgs:
            if msg.is_bot:
                display, initial, username = 'System', 'S', None
            elif msg.sender:
                display = (msg.sender.displayname if hasattr(msg.sender, 'displayname') and msg.sender.displayname
                           else msg.sender.username)
                initial = display[0].upper() if display else '?'
                username = msg.sender.username
            else:
                display, initial, username = 'System', 'S', None
            data.append({
                'id': msg.id,
                'text': msg.text,
                'sender_username': username,
                'sender_display': display,
                'sender_initial': initial,
                'created_at': msg.created_at.strftime('%H:%M'),
                'is_bot': msg.is_bot,
                'link': msg.link,
            })
        return JsonResponse({'status': 'success', 'messages': data, 'current_user': request.user.username})
    except Incident.DoesNotExist:
        return JsonResponse({'error': 'Incident not found'}, status=404)


@login_required(login_url='login')
@user_is_incident_responder
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def add_ioc(request, id):
    try:
        if request.method == 'POST':
            incident = Incident.objects.get(pk=id)
            ioc_type = request.POST.get('type')
            ioc_value = request.POST.get('value')
            ioc_description = request.POST.get('description', '')
            ioc_tags = request.POST.get('tag', '')

            # Create new IOC instance
            ioc = GenericIoc.objects.create(
                incident=incident,
                type=ioc_type,
                value=ioc_value,
                description=ioc_description,
                tag=ioc_tags
            )
            if ioc_type in THREAT_INTEL_ELIGIBLE_TYPES:
                schedule_lookup(ioc.id)
            Message.objects.create(
                incident=incident,
                sender=request.user,
                text=f"{request.user.username} added a new {ioc.get_type_display()} IoC: {ioc.value}",
                is_bot=True,
                link=f"/incident/{incident.id}/iocs"
            )
            _audit(incident, request, 'CREATE', 'IoC', f'Added {ioc.get_type_display()} IoC: {ioc.value}')
            supported = _cti_supported_types()
            if supported and ioc.type in supported:
                return redirect(f'/incident/{id}/iocs?check_new={ioc.id}')
            return redirect('/incident/' + str(id) + '/iocs')
        
    except Incident.DoesNotExist:
        return JsonResponse({'error': 'Invalid incident'}, status=400)
    except:
        return JsonResponse({'error': 'Invalid method'}, status=400)

@login_required(login_url='login')
@user_is_incident_responder
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def add_action(request, id):
    try:
        if request.method == 'POST':
            incident = Incident.objects.get(pk=id)
            data = json.loads(request.body)

            action_title =  data.get('title')
            action_description =  data.get('description')
            action_type =  data.get('type')
            action_iocs =  data.get('iocs', '')
            action_tags = data.get('tags', '')
            action_observed_at = data.get('observed_at')
            action_starting_time = data.get('starting_time')
            action_ending_time = data.get('ending_time')

            # Validate timing BEFORE creating anything, otherwise a bad request
            # leaves an orphan action behind while still returning an error.
            if not action_observed_at and not (action_starting_time and action_ending_time):
                return JsonResponse(
                    {'error': 'Please provide an Observed time, or both a Starting and Ending time.'},
                    status=400,
                )

            action = Action.objects.create(
                incident = incident,
                title = action_title,
                description = action_description,
                type = action_type,
                created_by = request.user,
                observed_at = action_observed_at or None,
                starting_time = None if action_observed_at else action_starting_time,
                ending_time = None if action_observed_at else action_ending_time,
            )

            if action_tags:
                action.tags.set(action_tags)
            if action_iocs:
                action.iocs.set(action_iocs)
            update_first_actions(incident=incident)
            sync_incident_times(incident)
            Message.objects.create(
                incident=incident,
                sender=request.user,
                text=f"{request.user.username} added timeline action: {action.title}",
                is_bot=True,
                link=f"/incident/{incident.id}/timeline"
            )
            _audit(incident, request, 'CREATE', 'Timeline', f'Added timeline event "{action.title}"')
            return JsonResponse({"status": "success", "message": "Action added successfully"})
        else:
            return JsonResponse({'error': 'Invalid method'}, status=400)
    except Incident.DoesNotExist:
        return JsonResponse({'error': 'Invalid incident'}, status=400)
    except Exception as e:
        return JsonResponse({'error': f'Other error :{e}'}, status=400)


# ── Bulk CSV import ────────────────────────────────────────────────────────
MAX_IMPORT_BYTES = 2 * 1024 * 1024  # 2 MB


def _read_csv_rows(request):
    """Return (rows, fieldmap, error_response). Shared CSV decoding for imports."""
    f = request.FILES.get('file')
    if not f:
        return None, None, JsonResponse({'error': 'No CSV file uploaded.'}, status=400)
    if f.size > MAX_IMPORT_BYTES:
        return None, None, JsonResponse({'error': 'File too large (max 2 MB).'}, status=400)
    try:
        text = f.read().decode('utf-8-sig')
    except (UnicodeDecodeError, AttributeError):
        return None, None, JsonResponse({'error': 'File must be a UTF-8 encoded CSV.'}, status=400)

    reader = csv.DictReader(StringIO(text))
    if not reader.fieldnames:
        return None, None, JsonResponse({'error': 'CSV is empty or has no header row.'}, status=400)
    fieldmap = {(h or '').strip().lower(): h for h in reader.fieldnames}
    return list(reader), fieldmap, None


def _import_parse_dt(s):
    """Parse an ISO-ish datetime string to an aware datetime, or None if blank/invalid."""
    from django.utils.dateparse import parse_datetime, parse_date
    s = (s or '').strip()
    if not s:
        return None
    dt = parse_datetime(s)
    if dt is None:
        d = parse_date(s)
        if d:
            dt = datetime(d.year, d.month, d.day)
    if dt is not None and timezone.is_naive(dt):
        dt = timezone.make_aware(dt, timezone.get_current_timezone())
    return dt


@login_required(login_url='login')
@user_is_incident_responder
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def import_iocs(request, id):
    """Bulk-create IoCs from an uploaded CSV (columns: type, value, description, status)."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid method'}, status=400)
    try:
        incident = Incident.objects.get(pk=id)
    except Incident.DoesNotExist:
        return JsonResponse({'error': 'Invalid incident'}, status=400)

    rows, fieldmap, err = _read_csv_rows(request)
    if err:
        return err
    if 'value' not in fieldmap:
        return JsonResponse({'error': "CSV must have a 'value' column. Expected columns: type, value, description, status."}, status=400)

    valid_types = {c[0] for c in GenericIoc._meta.get_field('type').choices}
    valid_status = {c[0] for c in GenericIoc._meta.get_field('status').choices}

    def cell(row, col):
        h = fieldmap.get(col)
        return (row.get(h) or '').strip() if h else ''

    created, skipped, errors = 0, 0, []
    for i, row in enumerate(rows, start=2):  # row 1 is the header
        value = cell(row, 'value')
        if not value:
            skipped += 1
            continue
        ioc_type = cell(row, 'type').upper()
        ioc_type = {'IPV4': 'IPADRESS', 'IPV6': 'IPADRESS', 'IP': 'IPADRESS'}.get(ioc_type, ioc_type)
        if ioc_type not in valid_types:
            ioc_type = 'OTHER'
        status = cell(row, 'status').upper()
        if status not in valid_status:
            status = 'SAFE'
        try:
            ioc = GenericIoc.objects.create(
                incident=incident,
                type=ioc_type,
                value=value,
                description=cell(row, 'description'),
                status=status,
                created_by=request.user,
                tag='',
            )
            created += 1
            if ioc.type in THREAT_INTEL_ELIGIBLE_TYPES:
                schedule_lookup(ioc.id)
        except Exception as e:
            errors.append(f'Row {i}: {e}')

    if created:
        Message.objects.create(
            incident=incident, sender=request.user, is_bot=True,
            text=f"{request.user.username} imported {created} IoC(s) from CSV",
            link=f"/incident/{incident.id}/iocs",
        )
        _audit(incident, request, 'CREATE', 'IoC', f'Bulk-imported {created} IoC(s) from CSV')

    return JsonResponse({'status': 'success', 'created': created, 'skipped': skipped, 'errors': errors[:20]})


@login_required(login_url='login')
@user_is_incident_responder
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def import_actions(request, id):
    """Bulk-create timeline actions from an uploaded CSV.

    Columns: title, description, type, observed_at (or starting_time + ending_time).
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid method'}, status=400)
    try:
        incident = Incident.objects.get(pk=id)
    except Incident.DoesNotExist:
        return JsonResponse({'error': 'Invalid incident'}, status=400)

    rows, fieldmap, err = _read_csv_rows(request)
    if err:
        return err
    if 'title' not in fieldmap:
        return JsonResponse({'error': "CSV must have a 'title' column. Expected columns: title, description, type, observed_at."}, status=400)

    valid_types = {c[0] for c in Action._meta.get_field('type').choices}

    def cell(row, col):
        h = fieldmap.get(col)
        return (row.get(h) or '').strip() if h else ''

    created, skipped, errors = 0, 0, []
    for i, row in enumerate(rows, start=2):
        title = cell(row, 'title')
        if not title:
            skipped += 1
            continue
        atype = cell(row, 'type').upper()
        if atype not in valid_types:
            atype = 'OTHER'
        observed_raw = cell(row, 'observed_at') or cell(row, 'time') or cell(row, 'date')
        starting_raw = cell(row, 'starting_time')
        ending_raw = cell(row, 'ending_time')
        if not observed_raw and not (starting_raw and ending_raw):
            errors.append(f'Row {i}: missing time (need observed_at, or both starting_time and ending_time)')
            skipped += 1
            continue
        try:
            if observed_raw:
                observed = _import_parse_dt(observed_raw)
                if observed is None:
                    raise ValueError(f"could not parse observed_at '{observed_raw}'")
                starting = ending = None
            else:
                observed = None
                starting = _import_parse_dt(starting_raw)
                ending = _import_parse_dt(ending_raw)
                if starting is None or ending is None:
                    raise ValueError('could not parse starting_time/ending_time')
            Action.objects.create(
                incident=incident,
                title=title[:200],
                description=cell(row, 'description'),
                type=atype,
                created_by=request.user,
                observed_at=observed,
                starting_time=starting,
                ending_time=ending,
            )
            created += 1
        except Exception as e:
            errors.append(f'Row {i}: {e}')

    if created:
        update_first_actions(incident=incident)
        sync_incident_times(incident)
        Message.objects.create(
            incident=incident, sender=request.user, is_bot=True,
            text=f"{request.user.username} imported {created} timeline action(s) from CSV",
            link=f"/incident/{incident.id}/timeline",
        )
        _audit(incident, request, 'CREATE', 'Timeline', f'Bulk-imported {created} timeline action(s) from CSV')

    return JsonResponse({'status': 'success', 'created': created, 'skipped': skipped, 'errors': errors[:20]})


@login_required(login_url='login')
@user_is_incident_responder
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def update_action(request, id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            action_id = data.get('id')
            
            if not action_id:
                return JsonResponse({'error': 'Action ID is required'}, status=400)

            action = Action.objects.get(pk=action_id)

            if action.incident.id != id:
                return JsonResponse({'error': 'Action is not associated with the specified incident'}, status=403)

            action.title = data.get('title')
            action.description = data.get('description')
            action.type = data.get('type')
            if data.get('observed_at'):
                action.observed_at = data.get('observed_at')
                action.starting_time = None
                action.ending_time = None
            else:
                action.starting_time = data.get('starting_time')
                action.ending_time = data.get('ending_time')
                action.observed_at = None
            action.iocs.set(data.get('iocs'))
            action.tags.set(data.get('tags'))
            
            action.save()
            update_first_actions(incident=action.incident)
            sync_incident_times(action.incident)
            _audit(action.incident, request, 'UPDATE', 'Timeline', f'Updated timeline event "{action.title}"')
            return JsonResponse({"status": "success", "message": "Action updated successfully"})

        except Action.DoesNotExist:
            return JsonResponse({'error': 'Action not found'}, status=404)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON format'}, status=400)
    return JsonResponse({'error': 'Invalid request method'}, status=405)

@login_required(login_url='login')
@user_is_incident_responder
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def delete_action(request, id):
    if request.method == 'DELETE':
        try:
            data = json.loads(request.body)
            action_id = data.get('id')
            
            if not action_id:
                return JsonResponse({'error': 'Action ID is required'}, status=400)

            action = Action.objects.get(pk=action_id)

            if action.incident.id != id:
                return JsonResponse({'error': 'Action is not associated with the specified incident'}, status=403)

            incident = action.incident
            action_title = action.title
            action.delete()
            update_first_actions(incident=incident)
            sync_incident_times(incident)
            _audit(incident, request, 'DELETE', 'Timeline', f'Deleted timeline event "{action_title}"')
            return JsonResponse({"status": "success", "message": "Action deleted successfully"})

        except Action.DoesNotExist:
            return JsonResponse({'error': 'Action not found'}, status=404)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON format'}, status=400)
    return JsonResponse({'error': 'Invalid request method'}, status=405)

@login_required(login_url='login')
@user_is_incident_responder_orpublic
def get_action(request, id, action_id):
    if request.method == 'GET':
        try:
            if not action_id:
                return JsonResponse({'error': 'Action ID is required'}, status=400)
            
            action = Action.objects.get(pk=action_id)
            if action.incident.id != id:
                return JsonResponse({'error': 'Action is not associated with the specified incident'}, status=403)

            action_data = {
                "id": action.id,
                "type": action.type,
                "title": action.title,
                "description": action.description,
                "created_at": action.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                "updated_at": action.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
                "observed_at": action.observed_at.strftime('%Y-%m-%d %H:%M:%S') if action.observed_at else "",
                "starting_time": action.starting_time.strftime('%Y-%m-%d %H:%M:%S') if action.starting_time else "",
                "ending_time": action.ending_time.strftime('%Y-%m-%d %H:%M:%S') if action.ending_time else "",
                "created_by": action.created_by.username if action.created_by else "Unknown",
                "tags": [{"id": tag.id, "name": tag.name, "color": tag.color} for tag in action.tags.all()],
                "iocs": [{"id": ioc.id, "value": ioc.value, "description": ioc.description, "type": ioc.type, "created_at": ioc.created_at.strftime('%Y-%m-%d %H:%M:%S')} for ioc in action.iocs.all()]
            }
            
            return JsonResponse({"status": "success", "data": action_data})

        except Action.DoesNotExist:
            return JsonResponse({'error': 'Action not found'}, status=404)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON format'}, status=400)
    return JsonResponse({'error': 'Invalid request method'}, status=405)



@login_required(login_url='login')
@user_is_incident_responder
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def add_task(request, id):
    try:
        if request.method == 'POST':
            data = json.loads(request.body)
            incident = Incident.objects.get(pk=id)
            task_title = data.get('title')
            task_description = data.get('description')
            task_status = data.get('status')
            task_priority = data.get('priority')
            task_assignee = data.get('assignee')
            task_external_reference = data.get('external_reference')
            task_tags = data.get('tags')
            task = Task.objects.create(
                incident=incident,
                title=task_title,
                description=task_description,
                status=task_status,
                priority=task_priority,
                external_reference=task_external_reference,
                created_by = request.user

            )
            if task_assignee:
                user = User.objects.filter(id=task_assignee).first()
                if user:
                    task.assignee = user
                    task.save()

            if task_tags:
                task.tags.set(task_tags)
                task.save()

            Message.objects.create(
                incident=incident,
                sender=request.user,
                text=f"{request.user.username} created task: {task.title}",
                is_bot=True,
                link=f"/incident/{incident.id}/tasks"
            )
            _audit(incident, request, 'CREATE', 'Task', f'Created task "{task.title}" [{task.priority} / {task.status}]')
            return redirect('/incident/' + str(id) + '/tasks')
        
    except Incident.DoesNotExist:
        return JsonResponse({'error': 'Invalid incident'}, status=400)
    except:
        return JsonResponse({'error': 'Invalid method'}, status=400)

@login_required(login_url='login')
@user_is_incident_responder
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def update_task(request, id):
    try:
        if request.method == 'POST':
            
            data = json.loads(request.body)
            incident = Incident.objects.get(pk=id)
            task_id = data.get('id')
            task = Task.objects.get(id=task_id)
            task.title = data.get('title')
            task.description = data.get('description')
            task.status = data.get('status')
            task.priority = data.get('priority')            
            task.external_reference = data.get('external_reference')
            task.tags.set(data.get('tags'))
            
            if data.get('assignee'):
                task.assignee = User.objects.get(pk=data.get('assignee'))
            else: 
                task.assignee = None
            task.save()
            _audit(incident, request, 'UPDATE', 'Task', f'Updated task "{task.title}"')
            return redirect('/incident/' + str(id) + '/tasks')

    except Incident.DoesNotExist:
        return JsonResponse({'error': 'Invalid incident'}, status=400)
    except:
        return JsonResponse({'error': 'Invalid method'}, status=400)

@login_required(login_url='login')
@user_is_incident_responder
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def delete_task(request, id):
    if request.method == 'DELETE':
        try:
            data = json.loads(request.body)
            task_id = data.get('id')
            
            if not task_id:
                return JsonResponse({'error': 'Task ID is required'}, status=400)

            task = Task.objects.get(pk=task_id)

            if task.incident.id != id:
                return JsonResponse({'error': 'Task is not associated with the specified incident'}, status=403)

            incident = task.incident
            task_title = task.title
            task.delete()
            _audit(incident, request, 'DELETE', 'Task', f'Deleted task "{task_title}"')
            return JsonResponse({"status": "success", "message": "Task deleted successfully"})

        except Task.DoesNotExist:
            return JsonResponse({'error': 'Task not found'}, status=404)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON format'}, status=400)
    return JsonResponse({'error': 'Invalid request method'}, status=405)


@login_required(login_url='login')
@user_is_incident_responder
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def delete_ioc(request, id):
    if request.method == 'DELETE':
        try:
            data = json.loads(request.body)
            ioc_id = data.get('ioc_id')
            
            if not ioc_id:
                return JsonResponse({'error': 'IoC ID is required'}, status=400)

            ioc = GenericIoc.objects.get(pk=ioc_id)

            if ioc.incident.id != id:
                return JsonResponse({'error': 'IoC is not associated with the specified incident'}, status=403)

            incident = ioc.incident
            ioc_value = ioc.value
            ioc.delete()
            _audit(incident, request, 'DELETE', 'IoC', f'Deleted IoC: {ioc_value}')
            return JsonResponse({"status": "success", "message": "IoC deleted successfully"})

        except GenericIoc.DoesNotExist:
            return JsonResponse({'error': 'IoC not found'}, status=404)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON format'}, status=400)
    return JsonResponse({'error': 'Invalid request method'}, status=405)


@login_required(login_url='login')
@user_is_incident_responder_orpublic
def get_all_iocs(request, id):
    if request.method == 'GET':
        try:
            incident = Incident.objects.get(pk=id)
            iocs = GenericIoc.objects.filter(incident=incident)
            
            ioc_list = [
                {
                    "id": ioc.id,
                    "value": ioc.value,
                    "type": ioc.type,
                    "description": ioc.description,
                    "status": ioc.status,
                    "reputation": ioc.reputation,
                }
                for ioc in iocs
            ]
            
            return JsonResponse({"status": "success", "iocs": ioc_list}, status=200)
            
        except Incident.DoesNotExist:
            return JsonResponse({'error': 'Incident not found'}, status=404)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid request method'}, status=405)


@login_required(login_url='login')
@user_is_incident_responder_orpublic
def get_ioc(request, id, ioc_id):
    if request.method == 'GET':
        try:
            if not ioc_id:
                return JsonResponse({'error': 'IoC ID is required'}, status=400)
            
            ioc = GenericIoc.objects.get(pk=ioc_id)
            if ioc.incident.id != id:
                return JsonResponse({'error': 'IoC is not associated with the specified incident'}, status=403)


            ioc_data = {
                "id": ioc.id,
                "value": ioc.value,
                "type": ioc.type,
                "description": ioc.description,
                "status": ioc.status,
                "reputation": ioc.reputation,
                "created_at": ioc.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                "updated_at": ioc.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
                "created_by": ioc.created_by.username if ioc.created_by else "Unknown",
                "tags": [{"id": tag.id, "name": tag.name, "color": tag.color} for tag in ioc.tags.all()],
                "actions": [{"id": action.id, "title": action.title, "description": action.description, "created_at": action.created_at.strftime('%Y-%m-%d %H:%M:%S')} for action in ioc.actions.all()]
            }
            
            return JsonResponse({"status": "success", "data": ioc_data})

        except GenericIoc.DoesNotExist:
            return JsonResponse({'error': 'IoC not found'}, status=404)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON format'}, status=400)
    return JsonResponse({'error': 'Invalid request method'}, status=405)

@login_required(login_url='login')
@user_is_incident_responder_orpublic
def get_impact(request, id, impact_id):
    if request.method == 'GET':
        try:
            if not impact_id:
                return JsonResponse({'error': 'Impact ID is required'}, status=400)
            
            impact = Impact.objects.get(pk=impact_id)
            if impact.incident.id != id:
                return JsonResponse({'error': 'Impact is not associated with the specified incident'}, status=403)


            impact_data = {
                "id": impact.id,
                "title": impact.title,
                "type": impact.type,
                "description": impact.description,
                "status": impact.status,
                "severity": impact.severity,
                "external_reference": impact.external_reference,
                "starting_time": impact.starting_time.strftime("%Y-%m-%dT%H:%M"),
                "ending_time": impact.ending_time.strftime("%Y-%m-%dT%H:%M"),
                "created_at": impact.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                "updated_at": impact.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
                "created_by": impact.created_by.username if impact.created_by else "Unknown",
                "tags": [{"id": tag.id, "name": tag.name, "color": tag.color} for tag in impact.tags.all()]
            }
            
            return JsonResponse({"status": "success", "data": impact_data})

        except Impact.DoesNotExist:
            return JsonResponse({'error': 'Impact not found'}, status=404)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON format'}, status=400)
    return JsonResponse({'error': 'Invalid request method'}, status=405)

        
@login_required(login_url='login')
@user_is_incident_responder
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def update_ioc(request, id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            ioc_id = data.get('ioc_id')
            ioc = GenericIoc.objects.get(id=ioc_id)

            ioc.type = data.get('type')
            ioc.value = data.get('value')
            ioc.description = data.get('description')
            ioc.save()
            _audit(ioc.incident, request, 'UPDATE', 'IoC', f'Updated IoC: {ioc.value}')
            return JsonResponse({'status': 'success', id: id})

        except GenericIoc.DoesNotExist:
            return JsonResponse({'error': 'Ioc not found'}, status=404)
    return JsonResponse({'error': 'Invalid request'}, status=400)

@login_required(login_url='login')
@user_is_incident_responder
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def update_impact(request, id):
    if request.method == 'POST':
        try:

            data = json.loads(request.body)['data']
            impact_id = data['id']
            impact = Impact.objects.get(id=impact_id)

            _IMPACT_EDITABLE = {'title', 'external_reference', 'description', 'status', 'severity', 'type', 'starting_time', 'ending_time'}
            for field in _IMPACT_EDITABLE:
                if field in data:
                    setattr(impact, field, data[field])

            impact.duration = datetime.strptime(impact.ending_time, "%Y-%m-%dT%H:%M") - datetime.strptime(impact.starting_time, "%Y-%m-%dT%H:%M")
            impact.save()
            
            return JsonResponse({'status': 'success', id: impact_id})

        except Impact.DoesNotExist as e:
            return JsonResponse({'error': f'Impact not found {e}'}, status=404)
    return JsonResponse({'error': 'Invalid request'}, status=400)


@login_required(login_url='login')
@user_is_incident_responder
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def delete_impact(request, id):
    if request.method == 'DELETE':
        try:
            data = json.loads(request.body)
            ioc_id = data.get('ioc_id')
            
            if not ioc_id:
                return JsonResponse({'error': 'IoC ID is required'}, status=400)

            ioc = GenericIoc.objects.get(pk=ioc_id)

            if ioc.incident.id != id:
                return JsonResponse({'error': 'IoC is not associated with the specified incident'}, status=403)

            ioc.delete()
            return JsonResponse({"status": "success", "message": "IoC deleted successfully"})

        except GenericIoc.DoesNotExist:
            return JsonResponse({'error': 'IoC not found'}, status=404)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON format'}, status=400)
    return JsonResponse({'error': 'Invalid request method'}, status=405)



@login_required(login_url='login')
@verify_permissions(['INCIDENT_LEAD'])
def update_role(request, id):
    if request.method != 'POST':
        return JsonResponse({"error": "Invalid request method"}, status=403)
    incident = get_object_or_404(Incident, id=id)
    body = json.loads(request.body.decode('utf-8'))
    user_id = body.get('user_id')

    new_role = body.get('role')
    display_role = body.get('display_role')
    
    if not user_id and not (new_role or display_role):
        return JsonResponse({"error": "Missing user_id or role or display_role"}, status=400)
    
    user = get_object_or_404(User, id=user_id)
    user_role, created = UserRole.objects.get_or_create(user=user, incident=incident)
    if new_role:
        if new_role not in dict(UserRole._meta.get_field('role').choices):
            return JsonResponse({"error": "Invalid role selected"}, status=400)
        user_role.role = new_role
        user_role.save()
        _audit(incident, request, 'UPDATE', 'Team', f'Changed role of {user.username} to {new_role}')
        return JsonResponse({"success": f"Role of {user.username} updated to {new_role}"}, status=200)

    if display_role:
        user_role.display_role = display_role
        user_role.save()
        _audit(incident, request, 'UPDATE', 'Team', f'Set display role of {user.username} to "{display_role}"')
        return JsonResponse({"success": f"Display Role of {user.username} updated to {display_role}"}, status=200)

@login_required(login_url='login')
@verify_permissions(['INCIDENT_LEAD'])
def delete_role(request, id):
    if request.method == 'POST':
        try:
            incident = get_object_or_404(Incident, id=id)

            body = json.loads(request.body.decode('utf-8'))
            user_id = body.get('user_id')
            userrole = UserRole.objects.get(incident=incident, user=user_id)

            # Check if the role being deleted is INCIDENT_LEAD
            if userrole.role == 'INCIDENT_LEAD':
                incident_id = userrole.incident.id
                
                # Count other INCIDENT_LEAD roles in the same incident
                incident_lead_count = UserRole.objects.filter(
                    incident_id=incident_id, 
                    role='INCIDENT_LEAD'
                ).exclude(id=userrole.id).count()
                
                if incident_lead_count == 0:
                    return JsonResponse({
                        'error': 'You must designate another Incident Lead before deleting this one.'
                    }, status=400)


            removed_user = userrole.user.username
            userrole.delete()
            _audit(incident, request, 'DELETE', 'Team', f'Removed responder {removed_user}')
            return JsonResponse({"status": "success", "message": "Responder deleted successfully"})
        except UserRole.DoesNotExist:
            return JsonResponse({'error': 'Responder not found'}, status=404)    
        
@login_required(login_url='login')
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def update_incident(request, id):
    if request.method == 'POST':
        try:
            # Fetch the incident object
            incident = Incident.objects.get(id=id)
            data = json.loads(request.body)

            # Verify the user role inside the update_incident function
            _ur = _get_user_role(request.user, incident)
            user_role = _ur.role if _ur else 'RESPONDER'

            if user_role == 'INCIDENT_LEAD':
                # If user is INCIDENT_LEAD, they can modify the title as well
                if 'title' in data:
                    incident.name = data['title']
            elif user_role == 'RESPONDER':
                # If user is RESPONDER, they can't modify the title, status, or severity
                if 'title' in data:
                    return JsonResponse({'error': 'Permission denied: You cannot modify the title'}, status=403)
                if 'status' in data:
                    return JsonResponse({'error': 'Permission denied: You cannot modify the status'}, status=403)
                if 'severity' in data:
                    return JsonResponse({'error': 'Permission denied: You cannot modify the severity'}, status=403)

            # Allow modification of other fields for both roles
            if 'description' in data:
                incident.description = data['description']
            if 'executive_summary' in data:
                incident.executive_summary = data['executive_summary']
            if 'lessons_learned' in data:
                incident.lessons_learned = data['lessons_learned']
            if 'technical_details' in data:
                incident.technical_details = data['technical_details']
            if 'external_reference' in data:
                incident.external_reference = (data['external_reference'] or '')[:255]
            if 'status' in data:
                new_status = data['status']
                valid_res = {r[0] for r in Incident.RESOLUTION_CHOICES}
                if new_status == 'CLOSED':
                    resolution = data.get('resolution', '').strip()
                    if not resolution:
                        resolution = incident.resolution  # keep existing if already set
                    if not resolution:
                        return JsonResponse({'error': 'resolution_required', 'message': 'A resolution is required to close an incident.'}, status=400)
                    if resolution not in valid_res:
                        return JsonResponse({'error': 'Invalid resolution.'}, status=400)
                    incident.resolution = resolution
                    incident.resolution_note = data.get('resolution_note', incident.resolution_note)[:2000]
                    sync_incident_times(incident)
                incident.status = new_status
            if 'resolution' in data and 'status' not in data:
                valid_res = {r[0] for r in Incident.RESOLUTION_CHOICES}
                if data['resolution'] in valid_res:
                    incident.resolution = data['resolution']
                    incident.resolution_note = data.get('resolution_note', '')[:2000]
            if 'severity' in data:
                incident.severity = data['severity']
            if 'export_include_timeline' in data:
                incident.export_include_timeline = data['export_include_timeline']
            if 'export_include_iocs' in data:
                incident.export_include_iocs = data['export_include_iocs']
            if 'export_include_attachements' in data:
                incident.export_include_attachements = data['export_include_attachements']
            if 'tlp' in data:
                valid_tlp = ('CLEAR', 'GREEN', 'AMBER', 'RED')
                if data['tlp'] in valid_tlp:
                    incident.tlp = data['tlp']
            if 'is_public' in data and user_role == 'INCIDENT_LEAD':
                incident.is_public = bool(data['is_public'])
            if 'ai_rephrase_enabled' in data and user_role == 'INCIDENT_LEAD':
                incident.ai_rephrase_enabled = bool(data['ai_rephrase_enabled'])

            # Save the updated incident
            incident.save()

            # Describe what changed for the audit log
            changed = [k for k in ('title','status','severity','tlp','is_public',
                                   'description','executive_summary','lessons_learned',
                                   'technical_details','external_reference') if k in data]
            if changed:
                _audit(incident, request, 'UPDATE', 'Settings',
                       f'Updated incident fields: {", ".join(changed)}')

            return JsonResponse({'status': 'success', 'message': 'Incident updated successfully'})
        except Incident.DoesNotExist:
            return JsonResponse({'error': 'Incident not found'}, status=404)
        

from django.contrib.auth import update_session_auth_hash
from django.contrib.auth.hashers import check_password

@login_required(login_url='login') 
def update_profile(request):
    if request.method == 'POST':
        try:
            user = request.user
            new_username = request.POST.get("username", user.username)

            # Check if the username is already taken by another user
            if User.objects.filter(username=new_username).exclude(id=user.id).exists():
                return JsonResponse({'error': "This username is already taken"}, status=403)

            user.displayname = request.POST.get("displayname", user.displayname)
            user.email = request.POST.get("email", user.email)
            user.light_mode = request.POST.get("light_mode")

            # Handle profile picture upload
            if "profile_picture" in request.FILES:
                profile_picture = request.FILES["profile_picture"]
                ext = os.path.splitext(profile_picture.name)[1].lower()
                if ext not in [".jpg", ".jpeg", ".png"]:
                    return JsonResponse({'error': "You ain't uploading any webshells on my app you dirty cow. Only JPG and PNG are allowed"}, status=403)
                try:
                    img = Image.open(io.BytesIO(profile_picture.read()))
                    img.verify()
                    profile_picture.seek(0)
                except Exception:
                    return JsonResponse({'error': 'Invalid image file.'}, status=400)

                user.profile_picture.save(profile_picture.name, profile_picture)

            # Handle password change
            current_password = request.POST.get("current_password")
            new_password = request.POST.get("new_password")
            confirm_password = request.POST.get("confirm_password")

            if current_password and new_password and confirm_password:
                if not check_password(current_password, user.password):
                    return JsonResponse({'error': "Current password is incorrect"}, status=403)
                
                if new_password != confirm_password:
                    return JsonResponse({'error': "New passwords do not match"}, status=403)

                user.set_password(new_password)  # Securely update password
                update_session_auth_hash(request, user)  # Keep the user logged in

            # Save notification preferences
            prefs = user.preferences if user.preferences else {}
            prefs['notify_assignment'] = 'notify_assignment' in request.POST
            prefs['notify_mention'] = 'notify_mention' in request.POST
            user.preferences = prefs

            # Save the updated user profile
            user.save()

            return render(request, 'profile.html', {
                'user': request.user,
                'prefs': prefs,
                'success': 'Profile updated successfully!',
            })

        except User.DoesNotExist:
            return JsonResponse({'error': 'User not found'}, status=404)


# ── Audit log API ─────────────────────────────────────────────────────────────

@login_required(login_url='login')
@user_is_incident_responder_orpublic
def get_audit_logs(request, id):
    """Return the last 500 audit log entries for an incident as JSON."""
    incident = get_object_or_404(Incident, pk=id)
    logs = (
        AuditLog.objects
        .filter(incident=incident)
        .select_related('user')
        .order_by('-timestamp')[:500]
    )
    data = [
        {
            'timestamp': log.timestamp.strftime('%Y-%m-%d %H:%M:%S'),
            'user':      log.user.username if log.user else '—',
            'ip':        log.ip_address or '—',
            'action':    log.action,
            'type':      log.object_type,
            'description': log.description,
        }
        for log in logs
    ]
    return JsonResponse({'logs': data})


@login_required(login_url='login')
@user_is_incident_responder_orpublic
def export_audit_logs(request, id):
    """Download audit logs as CSV."""
    incident = get_object_or_404(Incident, pk=id)
    logs = (
        AuditLog.objects
        .filter(incident=incident)
        .select_related('user')
        .order_by('-timestamp')
    )
    buf = StringIO()
    writer = csv.writer(buf)
    writer.writerow(['Timestamp (UTC)', 'User', 'IP Address', 'Action', 'Type', 'Description'])
    for log in logs:
        writer.writerow([
            log.timestamp.strftime('%Y-%m-%d %H:%M:%S'),
            log.user.username if log.user else '',
            log.ip_address,
            log.action,
            log.object_type,
            log.description,
        ])
    _audit(incident, request, 'EXPORT', 'Audit Log', 'Exported audit log as CSV')
    response = HttpResponse(buf.getvalue(), content_type='text/csv')
    response['Content-Disposition'] = (
        f'attachment; filename="audit_log_incident_{id}.csv"'
    )
    return response


# ─────────────────────────────────────────────────────────────────────────────
# AI REPHRASE
# ─────────────────────────────────────────────────────────────────────────────

def _ai_rephrase_rate_limit_exceeded(user_id: int) -> bool:
    """Return True if this user has used up their daily AI rephrase quota."""
    from datetime import timezone as dt_timezone
    now_utc = timezone.now()  # always UTC-aware in Django with USE_TZ=True
    today_utc = now_utc.astimezone(dt_timezone.utc).date().isoformat()
    key = f'ai_rephrase_{user_id}_{today_utc}'
    try:
        count = cache.incr(key)
    except ValueError:
        midnight_utc = (now_utc.astimezone(dt_timezone.utc) + timedelta(days=1)).replace(
            hour=0, minute=0, second=0, microsecond=0
        )
        secs = int((midnight_utc - now_utc).total_seconds())
        cache.set(key, 1, timeout=secs)
        count = 1
    return count > _AI_REPHRASE_DAILY_LIMIT


@login_required(login_url='login')
@user_is_incident_responder_orpublic
def ai_rephrase(request, id):
    """
    POST  /api/incident/<id>/ai-rephrase/
    Body: { "field": "description" | "executive_summary" | "technical_details" | "lessons_learned" }

    Generates text for the requested field using Anthropic Claude (preferred)
    or OpenAI GPT depending on which API key is configured on the server.
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)

    incident = get_object_or_404(Incident, pk=id)

    if not incident.ai_rephrase_enabled:
        return JsonResponse({'error': 'AI rephrase is not enabled for this incident.'}, status=403)

    # Restrict to write-capable roles
    _ur = _get_user_role(request.user, incident)
    if _ur and _ur.role in ('READER', 'PUBLIC_VIEWER'):
        return JsonResponse({'error': 'Permission denied: read-only role.'}, status=403)

    if _ai_rephrase_rate_limit_exceeded(request.user.pk):
        return JsonResponse(
            {'error': f'Daily AI rephrase limit ({_AI_REPHRASE_DAILY_LIMIT} calls) reached. Resets at midnight UTC.'},
            status=429,
        )

    try:
        body = json.loads(request.body)
        field = body.get('field', '').strip()
    except (json.JSONDecodeError, AttributeError):
        return JsonResponse({'error': 'Invalid JSON body.'}, status=400)

    valid_fields = ('description', 'executive_summary', 'technical_details', 'lessons_learned')
    if field not in valid_fields:
        return JsonResponse(
            {'error': f'Invalid field. Must be one of: {", ".join(valid_fields)}'},
            status=400,
        )

    # Resolve provider and API key — DB takes priority, env vars are fallback
    anthropic_key = ''
    openai_key    = ''
    ps = PlatformSettings.get()

    if ps.ai_provider == 'ANTHROPIC' and ps.ai_api_key:
        anthropic_key = ps.ai_api_key
    elif ps.ai_provider == 'OPENAI' and ps.ai_api_key:
        openai_key = ps.ai_api_key
    else:
        # Fallback: environment / settings.py (dev convenience)
        from django.conf import settings as _settings
        _ant = getattr(_settings, 'ANTHROPIC_API_KEY', '') or os.environ.get('ANTHROPIC_API_KEY', '')
        _oai = getattr(_settings, 'OPENAI_API_KEY',    '') or os.environ.get('OPENAI_API_KEY',    '')
        if _ant:
            anthropic_key = _ant
        elif _oai:
            openai_key = _oai

    if not anthropic_key and not openai_key:
        return JsonResponse(
            {'error': 'AI is not configured. An administrator must set the AI provider '
                      'and API key in Platform Settings.'},
            status=503,
        )

    # ── Build incident context ────────────────────────────────────────────
    actions = incident.actions.all().order_by('observed_at', 'starting_time')
    iocs    = incident.genericiocs.all()
    tasks   = incident.tasks.all()

    def _fmt(dt):
        return dt.strftime('%Y-%m-%d %H:%M') if dt else 'unknown'

    action_lines = [
        f"  [{_fmt(a.observed_at or a.starting_time)}] {a.type}: {a.title} — {a.description or '(no description)'}"
        for a in actions
    ] or ['  (no events recorded)']

    ioc_lines = [
        f"  {ioc.type}: {ioc.value} — {ioc.description or '(no description)'}"
        for ioc in iocs
    ] or ['  (no IoCs recorded)']

    task_lines = [
        f"  [{t.status}] {t.title}: {t.description or '(no description)'}"
        for t in tasks
    ] or ['  (no tasks recorded)']

    FIELD_INSTRUCTIONS = {
        'description': (
            'Write a concise incident description (2–4 sentences). '
            'Summarise what happened: incident type, initial vector if known, affected systems, and scope. '
            'Use factual, neutral language. Plain text only — no markdown, no headings.'
        ),
        'executive_summary': (
            'Write a professional executive summary (3–5 sentences) for a non-technical audience. '
            'Focus on business impact, what was affected, the response taken, and current status. '
            'Avoid technical jargon. Plain text only — no markdown, no headings.'
        ),
        'technical_details': (
            'Write a detailed technical analysis. '
            'Cover the attack vector, techniques observed (reference specific IoCs and timeline events), '
            'affected systems, and containment actions taken. '
            'Be precise and technical. Plain text only — no markdown, no headings.'
        ),
        'lessons_learned': (
            'Write a lessons learned section covering: what detection gaps existed, what response actions worked well, '
            'what could have been improved, and 3–5 concrete recommendations to prevent recurrence. '
            'Be actionable. Plain text only — no markdown, no headings.'
        ),
    }

    field_display = field.replace('_', ' ').title()

    prompt = (
        f'You are a senior incident response analyst writing a professional security incident report.\n\n'
        f'Based on the incident data below, write ONLY the "{field_display}" section.\n'
        f'Do NOT include a heading or section title. Output plain prose only — no markdown, '
        f'no bullet points, no bold/italic formatting.\n\n'
        f'INCIDENT: {incident.name}\n'
        f'STATUS: {incident.status} | SEVERITY: {incident.severity} | TLP: {incident.tlp}\n'
        f'STARTED: {_fmt(incident.starting_time)} | '
        f'ENDED: {_fmt(incident.ending_time) if incident.ending_time else "ongoing"}\n\n'
        f'TIMELINE ({actions.count()} events):\n' + '\n'.join(action_lines) + '\n\n'
        f'INDICATORS OF COMPROMISE ({iocs.count()} IoCs):\n' + '\n'.join(ioc_lines) + '\n\n'
        f'TASKS ({tasks.count()} tasks):\n' + '\n'.join(task_lines) + '\n\n'
        f'INSTRUCTION: {FIELD_INSTRUCTIONS[field]}'
    )

    # ── Call AI provider ──────────────────────────────────────────────────
    try:
        if anthropic_key:
            import anthropic as _anthropic_sdk
            client  = _anthropic_sdk.Anthropic(api_key=anthropic_key)
            message = client.messages.create(
                model='claude-3-5-haiku-20241022',
                max_tokens=1024,
                messages=[{'role': 'user', 'content': prompt}],
            )
            generated = message.content[0].text.strip()
            provider  = 'Anthropic'
        else:
            from openai import OpenAI as _OpenAI
            client   = _OpenAI(api_key=openai_key)
            response = client.chat.completions.create(
                model='gpt-4o-mini',
                max_tokens=1024,
                messages=[{'role': 'user', 'content': prompt}],
            )
            generated = response.choices[0].message.content.strip()
            provider  = 'OpenAI'

        _audit(incident, request, 'UPDATE', 'AI Rephrase',
               f'AI-generated "{field_display}" via {provider}')
        return JsonResponse({'text': generated, 'field': field, 'provider': provider})

    except Exception as exc:
        return JsonResponse({'error': f'AI generation failed: {exc}'}, status=500)


# ─────────────────────────────────────────────────────────────────────────────
# THREAT INTEL
# ─────────────────────────────────────────────────────────────────────────────

@login_required(login_url='login')
@user_is_incident_responder_orpublic
def refresh_ioc_reputation(request, id, ioc_id):
    """POST — re-trigger threat intel lookup for an IoC."""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    ioc = get_object_or_404(GenericIoc, pk=ioc_id, incident_id=id)
    if ioc.type in THREAT_INTEL_ELIGIBLE_TYPES:
        schedule_lookup(ioc.id)
        return JsonResponse({'status': 'queued'})
    return JsonResponse({'status': 'skipped', 'reason': 'type not eligible'})


# ─────────────────────────────────────────────────────────────────────────────
# WAR ROOM
# ─────────────────────────────────────────────────────────────────────────────

@login_required(login_url='login')
@user_is_incident_responder_orpublic
def warroom(request, id):
    incident = get_object_or_404(Incident, pk=id)
    user_role = _get_user_role(request.user, incident)
    if user_role is None:
        return _forbidden(request, 'You are not a member of this incident.', incident=incident)
    return render(request, 'incidents/warroom.html', {
        'incident': incident,
        'user': request.user,
        'current_user_role': user_role,
    })


@login_required(login_url='login')
@user_is_incident_responder_orpublic
def warroom_data(request, id):
    """GET — JSON snapshot for the war room auto-refresh."""
    incident = get_object_or_404(Incident, pk=id)

    # Timeline — last 10 events
    actions_qs = (
        incident.actions.all()
        .order_by('-observed_at', '-starting_time')
        .select_related('created_by')[:10]
    )
    timeline = []
    for a in actions_qs:
        ts = a.observed_at or a.starting_time
        timeline.append({
            'id': a.id,
            'type': a.type,
            'title': a.title,
            'description': a.description,
            'ts': ts.strftime('%d %b %H:%M') if ts else '',
            'created_by': a.created_by.username if a.created_by else '',
        })

    # Messages — last 10
    msgs_qs = (
        incident.messages.all()
        .order_by('-created_at')
        .select_related('sender')[:10]
    )
    messages_data = []
    for m in reversed(list(msgs_qs)):
        if m.is_bot:
            sender = 'System'
        elif m.sender:
            sender = m.sender.displayname or m.sender.username
        else:
            sender = 'System'
        messages_data.append({
            'id': m.id,
            'text': m.text,
            'sender': sender,
            'is_bot': m.is_bot,
            'ts': m.created_at.strftime('%H:%M'),
        })

    # KPIs
    total_iocs   = incident.genericiocs.count()
    tasks_all    = list(incident.tasks.values('id', 'status'))
    open_tasks   = sum(1 for t in tasks_all if t['status'] in ('OPEN', 'IN_PROGRESS'))
    total_tasks  = len(tasks_all)
    responders   = list(
        incident.incident_roles.select_related('user')
        .exclude(role='PUBLIC_VIEWER')
    )

    # Open tasks detail
    open_tasks_qs = (
        incident.tasks.filter(status__in=('OPEN', 'IN_PROGRESS'))
        .select_related('assignee')
        .order_by('priority', 'created_at')[:20]
    )
    open_tasks_list = [
        {
            'id': t.id,
            'title': t.title,
            'status': t.status,
            'priority': t.priority,
            'assignee': t.assignee.username if t.assignee else None,
        }
        for t in open_tasks_qs
    ]

    # Critical IoCs (reputation malicious)
    critical_iocs = [
        {
            'id': ioc.id,
            'type': ioc.type,
            'value': ioc.value,
            'reputation': ioc.reputation,
        }
        for ioc in incident.genericiocs.all()
        if ioc.reputation and ioc.reputation.get('status') == 'malicious'
    ]

    # Responders list
    responders_list = [
        {
            'username': ur.user.username,
            'display': ur.user.displayname or ur.user.username,
            'role': ur.role,
            'display_role': ur.display_role or ur.get_role_display(),
        }
        for ur in responders
    ]

    return JsonResponse({
        'incident': {
            'id': incident.id,
            'name': incident.name,
            'severity': incident.severity,
            'status': incident.status,
            'starting_time': incident.starting_time.isoformat() if incident.starting_time else None,
        },
        'timeline': timeline,
        'messages': messages_data,
        'kpis': {
            'total_iocs': total_iocs,
            'open_tasks': open_tasks,
            'total_tasks': total_tasks,
            'responders': len(responders),
        },
        'open_tasks': open_tasks_list,
        'critical_iocs': critical_iocs,
        'responders': responders_list,
    })


# ─────────────────────────────────────────────────────────────────────────────
# PLATFORM SETTINGS (admin-only)
# ─────────────────────────────────────────────────────────────────────────────
# INCIDENT CATEGORIES
# ─────────────────────────────────────────────────────────────────────────────

@login_required(login_url='login')
def list_categories(request):
    cats = list(IncidentCategory.objects.order_by('name').values('id', 'name', 'color'))
    return JsonResponse({'categories': cats})


@login_required(login_url='login')
def create_category(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid method'}, status=405)
    if not request.user.is_superuser and not UserRole.objects.filter(
        user=request.user, role__in=['INCIDENT_LEAD', 'RESPONDER']
    ).exists():
        return JsonResponse({'error': 'Forbidden'}, status=403)
    data = json.loads(request.body)
    name = data.get('name', '').strip()
    color = data.get('color', '#796FA7').strip()
    if not name:
        return JsonResponse({'error': 'Name is required'}, status=400)
    cat, created = IncidentCategory.objects.get_or_create(name=name, defaults={'color': color})
    return JsonResponse({'id': cat.id, 'name': cat.name, 'color': cat.color, 'created': created})


@login_required(login_url='login')
@user_is_incident_responder
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def add_incident_category(request, id):
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid method'}, status=405)
    incident = get_object_or_404(Incident, id=id)
    data = json.loads(request.body)
    category_id = data.get('category_id')
    if not category_id:
        return JsonResponse({'error': 'category_id required'}, status=400)
    cat = get_object_or_404(IncidentCategory, id=category_id)
    incident.categories.add(cat)
    return JsonResponse({'status': 'success', 'id': cat.id, 'name': cat.name, 'color': cat.color})


@login_required(login_url='login')
@user_is_incident_responder
@verify_permissions(['INCIDENT_LEAD', 'RESPONDER'])
def remove_incident_category(request, id):
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid method'}, status=405)
    incident = get_object_or_404(Incident, id=id)
    data = json.loads(request.body)
    category_id = data.get('category_id')
    if not category_id:
        return JsonResponse({'error': 'category_id required'}, status=400)
    incident.categories.remove(category_id)
    return JsonResponse({'status': 'success'})


# ─────────────────────────────────────────────────────────────────────────────

@login_required(login_url='login')
def save_platform_settings(request):
    """
    POST  /api/platform-settings/
    Admin-only. Saves platform-wide AI provider / API key.
    Body: { "ai_provider": "ANTHROPIC"|"OPENAI"|"NONE", "ai_api_key": "..." }
    Leave ai_api_key empty (or omit) to keep the existing key unchanged.
    """
    if not request.user.is_admin:
        return JsonResponse({'error': 'Admin access required.'}, status=403)
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed.'}, status=405)

    try:
        body = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON.'}, status=400)

    valid_providers = ('NONE', 'ANTHROPIC', 'OPENAI')
    provider = body.get('ai_provider', '').upper()
    if provider not in valid_providers:
        return JsonResponse({'error': f'Invalid provider. Must be one of: {valid_providers}'}, status=400)

    ps = PlatformSettings.get()
    ps.ai_provider = provider

    new_key = body.get('ai_api_key', '').strip()
    if new_key:                # only overwrite if a value was submitted
        ps.ai_api_key = new_key

    if provider == 'NONE':     # if disabled, clear the stored key too
        ps.ai_api_key = ''

    ps.save()
    return JsonResponse({
        'status':      'saved',
        'ai_provider': ps.ai_provider,
        'key_set':     bool(ps.ai_api_key),
    })


# ─────────────────────────────────────────────────────────────────────────────
# CTI PROVIDERS (admin-only management)
# ─────────────────────────────────────────────────────────────────────────────

@login_required(login_url='login')
def save_cti_provider(request):
    """
    POST /api/cti-providers/
    Upsert (create or update) a CTI provider for the platform.
    Body: { name, api_key, base_url, enabled }
    Leaving api_key blank keeps the existing key.
    """
    if not request.user.is_admin:
        return JsonResponse({'error': 'Admin access required.'}, status=403)
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed.'}, status=405)

    try:
        body = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON.'}, status=400)

    name = body.get('name', '').upper()
    valid = [c[0] for c in CtiProvider.PROVIDER_CHOICES]
    if name not in valid:
        return JsonResponse({'error': f'Unknown provider. Valid: {valid}'}, status=400)

    provider, _ = CtiProvider.objects.get_or_create(name=name)

    new_key = body.get('api_key', '').strip()
    if new_key:
        provider.api_key = new_key

    base_url = body.get('base_url', '').strip()
    if base_url:
        provider.base_url = base_url

    if 'enabled' in body:
        provider.enabled = bool(body['enabled'])

    provider.save()
    return JsonResponse({
        'status':  'saved',
        'name':    provider.name,
        'enabled': provider.enabled,
        'key_set': bool(provider.api_key),
    })


@login_required(login_url='login')
def delete_cti_provider(request):
    """POST /api/cti-providers/delete/  Body: { name }"""
    if not request.user.is_admin:
        return JsonResponse({'error': 'Admin access required.'}, status=403)
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed.'}, status=405)

    try:
        body = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON.'}, status=400)

    name = body.get('name', '').upper()
    CtiProvider.objects.filter(name=name).delete()
    return JsonResponse({'status': 'deleted', 'name': name})


# ─────────────────────────────────────────────────────────────────────────────
# IOC REPUTATION LOOKUP
# ─────────────────────────────────────────────────────────────────────────────

def _cti_supported_types():
    """Return the set of GenericIoc.type values that have at least one enabled CTI integration."""
    types = set()
    if CtiProvider.objects.filter(name='VIRUSTOTAL', enabled=True).exclude(api_key='').exists():
        types.update({'IPADRESS', 'DOMAIN', 'HASH', 'FILE', 'URL'})
    # Future: AbuseIPDB / Shodan → 'IPADRESS'
    return types


def _vt_lookup(ioc, api_key):
    """Query VirusTotal v3 for a GenericIoc. Returns a result dict or None."""
    import httpx, base64

    base    = 'https://www.virustotal.com/api/v3'
    headers = {'x-apikey': api_key}
    itype   = ioc.type
    value   = ioc.value.strip()

    if itype == 'IPADRESS':
        endpoint = f'{base}/ip_addresses/{value}'
        gui_path = f'ip-address/{value}'
    elif itype == 'DOMAIN':
        endpoint = f'{base}/domains/{value}'
        gui_path = f'domain/{value}'
    elif itype in ('HASH', 'FILE'):
        endpoint = f'{base}/files/{value}'
        gui_path = f'file/{value}'
    elif itype == 'URL':
        url_id   = base64.urlsafe_b64encode(value.encode()).decode().rstrip('=')
        endpoint = f'{base}/urls/{url_id}'
        gui_path = f'url/{url_id}'
    else:
        return None   # type not supported by VT

    try:
        resp = httpx.get(endpoint, headers=headers, timeout=10)
    except Exception as exc:
        return {'provider': 'VirusTotal', 'verdict': 'error', 'detail': str(exc), 'link': None}

    if resp.status_code == 404:
        return {'provider': 'VirusTotal', 'verdict': 'unknown',
                'detail': 'Not found in VirusTotal database', 'link': None}
    if resp.status_code == 401:
        return {'provider': 'VirusTotal', 'verdict': 'error',
                'detail': 'Invalid API key', 'link': None}
    if resp.status_code != 200:
        return {'provider': 'VirusTotal', 'verdict': 'error',
                'detail': f'HTTP {resp.status_code}', 'link': None}

    attrs      = resp.json().get('data', {}).get('attributes', {})
    stats      = attrs.get('last_analysis_stats', {})
    malicious  = stats.get('malicious',  0)
    suspicious = stats.get('suspicious', 0)
    harmless   = stats.get('harmless',   0)
    undetected = stats.get('undetected', 0)
    total      = sum(stats.values())

    if malicious > 0:
        verdict = 'malicious'
    elif suspicious > 0:
        verdict = 'suspicious'
    elif total == 0:
        verdict = 'unknown'
    else:
        verdict = 'clean'

    result = {
        'provider':   'VirusTotal',
        'verdict':    verdict,
        'score':      f'{malicious}/{total}',
        'malicious':  malicious,
        'suspicious': suspicious,
        'harmless':   harmless,
        'undetected': undetected,
        'total':      total,
        'link':       f'https://www.virustotal.com/gui/{gui_path}',
    }
    # Extra contextual fields (IP/domain/file)
    for field in ('country', 'asn', 'as_owner', 'meaningful_name', 'type_description',
                  'network', 'last_analysis_date'):
        val = attrs.get(field)
        if val is not None:
            result[field] = val
    return result


@login_required(login_url='login')
@user_is_incident_responder_orpublic
def ioc_reputation(request, id, ioc_id):
    """
    GET /api/incident/<id>/ioc/<ioc_id>/reputation/
    Query all enabled CTI providers, persist results, return JSON for the UI.
    Response: { status: 'ok', reputation: { status, checked_at, vt: {...} } }
    """
    from django.utils import timezone as tz
    incident = get_object_or_404(Incident, pk=id)
    ioc      = get_object_or_404(GenericIoc, pk=ioc_id, incident=incident)

    providers = CtiProvider.objects.filter(enabled=True).exclude(api_key='')
    if not providers.exists():
        return JsonResponse({'status': 'error', 'error': 'No CTI providers configured.'}, status=503)

    raw_results = []
    for p in providers:
        if p.name == 'VIRUSTOTAL':
            r = _vt_lookup(ioc, p.api_key)
            if r:
                raw_results.append(r)
        # Future providers: ABUSEIPDB, SHODAN, OTXALIENVAULT, MISP …

    if not raw_results:
        return JsonResponse({
            'status': 'error',
            'error':  'No configured provider supports this IoC type.',
        }, status=422)

    # Determine worst verdict across all providers
    SEVERITY = {'malicious': 3, 'suspicious': 2, 'unknown': 1, 'clean': 0, 'error': -1}
    worst_verdict = max(
        (r.get('verdict', 'error') for r in raw_results),
        key=lambda v: SEVERITY.get(v, -1),
    )
    # Map 'error' → 'unknown' for display
    status = worst_verdict if worst_verdict in SEVERITY and worst_verdict != 'error' else 'unknown'

    # Build reputation object in the shape renderRepModal expects
    reputation = {
        'status':     status,
        'checked_at': tz.now().isoformat(),
    }
    for r in raw_results:
        if r.get('provider') == 'VirusTotal':
            reputation['vt'] = {k: v for k, v in r.items()
                                if k not in ('provider', 'verdict', 'score')}

    # Persist to the ioc record
    ioc.reputation = reputation
    ioc.save(update_fields=['reputation'])
    _audit(incident, request, 'UPDATE', 'IoC', f'Reputation checked for IoC #{ioc.id} ({ioc.value}) — verdict: {status}')

    return JsonResponse({'status': 'ok', 'reputation': reputation})


# ─────────────────────────────────────────────────────────────────────────────
# THREAT INTELLIGENCE HUB
# ─────────────────────────────────────────────────────────────────────────────

def _accessible_incidents(user):
    """Return QS of incidents the user can read (any role, or public, or platform role)."""
    from django.db.models import Q
    pr = getattr(user, 'platform_role', '') or ''
    if user.is_superuser or pr in ('SOC_ANALYST', 'SOC_LEAD'):
        return Incident.objects.all()
    role_incident_ids = UserRole.objects.filter(user=user).values_list('incident_id', flat=True)
    return Incident.objects.filter(
        Q(id__in=role_incident_ids) | Q(is_public=True)
    )


def _get_user_role(user, incident):
    """
    Return the effective UserRole for a user on an incident.
    For platform-role users with no explicit incident role, returns a synthetic
    (unsaved) UserRole so callers can check .role uniformly.
    Returns None if the user has no access at all.
    """
    try:
        return UserRole.objects.get(user=user, incident=incident)
    except UserRole.DoesNotExist:
        pr = getattr(user, 'platform_role', '') or ''
        if user.is_superuser or pr == 'SOC_LEAD':
            return UserRole(user=user, incident=incident, role='INCIDENT_LEAD')
        if pr == 'SOC_ANALYST':
            return UserRole(user=user, incident=incident, role='RESPONDER')
        if incident.is_public:
            return UserRole(user=user, incident=incident, role='PUBLIC_VIEWER')
        return None


def _ti_ioc_queryset(user, params):
    """
    Build a filtered GenericIoc queryset scoped to accessible incidents.
    params is a dict-like (request.GET).
    """
    from django.db.models import Q
    incidents = _accessible_incidents(user)

    date_from = params.get('date_from', '')
    date_to   = params.get('date_to',   '')
    ioc_type  = params.get('type',      '')
    ioc_status = params.get('status',   '')
    verdict   = params.get('verdict',   '')
    tlp       = params.get('tlp',       '')
    campaign_id = params.get('campaign', '')
    search    = params.get('search',    '').strip()
    incident_id = params.get('incident', '')

    if campaign_id:
        try:
            campaign = Campaign.objects.get(pk=int(campaign_id))
            incidents = incidents.filter(id__in=campaign.incidents.values_list('id', flat=True))
        except (Campaign.DoesNotExist, ValueError):
            pass

    if incident_id:
        try:
            incidents = incidents.filter(pk=int(incident_id))
        except ValueError:
            pass

    qs = GenericIoc.objects.filter(incident__in=incidents).select_related('incident', 'created_by')

    if date_from:
        try:
            from datetime import date
            qs = qs.filter(created_at__date__gte=date.fromisoformat(date_from))
        except ValueError:
            pass
    if date_to:
        try:
            from datetime import date
            qs = qs.filter(created_at__date__lte=date.fromisoformat(date_to))
        except ValueError:
            pass
    if ioc_type:
        qs = qs.filter(type=ioc_type)
    if ioc_status:
        qs = qs.filter(status=ioc_status)
    if verdict:
        # Filter on the JSON reputation.status field
        qs = qs.filter(reputation__status=verdict)
    if tlp:
        VALID_TLP = {'CLEAR', 'GREEN', 'AMBER', 'RED'}
        if tlp.upper() in VALID_TLP:
            qs = qs.filter(incident__tlp=tlp.upper())
    if search:
        qs = qs.filter(value__icontains=search)

    return qs


@login_required(login_url='login')
def threat_intel(request):
    """GET /threat-intel/ — Threat Intelligence Hub page."""
    from django.db.models import Count

    incidents   = _accessible_incidents(request.user)
    campaigns   = Campaign.objects.all().prefetch_related('incidents')
    all_iocs    = GenericIoc.objects.filter(incident__in=incidents)
    total_iocs  = all_iocs.count()
    unique_iocs = all_iocs.values('value').distinct().count()

    malicious_count = all_iocs.filter(reputation__status='malicious').count()
    pct_malicious   = round(malicious_count / total_iocs * 100, 1) if total_iocs else 0

    from .choices_processor import choices_context
    from django.http import HttpRequest
    req = HttpRequest()
    choices = choices_context(req)

    context = {
        'incidents':       incidents.order_by('-created_at'),
        'campaigns':       campaigns,
        'total_iocs':      total_iocs,
        'unique_iocs':     unique_iocs,
        'pct_malicious':   pct_malicious,
        'incident_count':  incidents.count(),
        'campaign_count':  campaigns.count(),
        'IOC_TYPE_CHOICES':   choices['GENERIC_IOC_TYPE_CHOICES'],
        'IOC_STATUS_CHOICES': choices['GENERIC_IOC_STATUS_CHOICES'],
        'cti_configured':  CtiProvider.objects.filter(enabled=True).exclude(api_key='').exists(),
    }
    return render(request, 'threat_intel.html', context)


@login_required(login_url='login')
def api_ti_iocs(request):
    """
    GET /api/threat-intel/iocs/
    Returns paginated IOC list with cross-incident occurrence counts.
    """
    from django.db.models import Count, Value
    from django.db.models.functions import Lower

    PAGE_SIZE = 50
    try:
        page = max(1, int(request.GET.get('page', 1)))
    except ValueError:
        page = 1

    qs = _ti_ioc_queryset(request.user, request.GET).order_by('-created_at')

    # Annotate with how many accessible incidents share the same value
    accessible_ids = list(_accessible_incidents(request.user).values_list('id', flat=True))

    total = qs.count()
    offset = (page - 1) * PAGE_SIZE
    iocs = qs[offset: offset + PAGE_SIZE]

    # Build value → occurrence count map for this page's values
    page_values = [i.value for i in iocs]
    occ_map = {}
    if page_values:
        rows = (
            GenericIoc.objects
            .filter(incident_id__in=accessible_ids, value__in=page_values)
            .values('value')
            .annotate(cnt=Count('id'))
        )
        occ_map = {r['value']: r['cnt'] for r in rows}

    data = []
    for ioc in iocs:
        rep = ioc.reputation or {}
        data.append({
            'id':          ioc.id,
            'value':       ioc.value,
            'type':        ioc.type,
            'status':      ioc.status,
            'verdict':     rep.get('status', 'unknown'),
            'vt_score':    rep.get('vt', {}).get('malicious', 0) if rep.get('vt') else None,
            'incident_id': ioc.incident_id,
            'incident':    ioc.incident.name,
            'created_at':  ioc.created_at.isoformat(),
            'occurrences': occ_map.get(ioc.value, 1),
        })

    return JsonResponse({
        'iocs':      data,
        'total':     total,
        'page':      page,
        'page_size': PAGE_SIZE,
        'pages':     max(1, (total + PAGE_SIZE - 1) // PAGE_SIZE),
    })


@login_required(login_url='login')
def api_ti_stats(request):
    """GET /api/threat-intel/stats/ — KPI numbers."""
    from django.db.models import Count

    qs = _ti_ioc_queryset(request.user, request.GET)
    total   = qs.count()
    unique  = qs.values('value').distinct().count()
    mal     = qs.filter(reputation__status='malicious').count()
    susp    = qs.filter(reputation__status='suspicious').count()
    clean   = qs.filter(reputation__status='clean').count()
    inc_cnt = qs.values('incident_id').distinct().count()

    by_type = list(
        qs.values('type').annotate(cnt=Count('id')).order_by('-cnt')
    )

    return JsonResponse({
        'total':           total,
        'unique':          unique,
        'malicious':       mal,
        'suspicious':      susp,
        'clean':           clean,
        'incident_count':  inc_cnt,
        'by_type':         by_type,
    })


@login_required(login_url='login')
def api_ti_heatmap(request):
    """
    GET /api/threat-intel/heatmap/
    Returns daily IOC counts for the calendar heatmap.
    Response: { days: [ {date, count}, … ] }  — last 365 days by default.
    """
    from django.db.models import Count
    from datetime import date, timedelta

    qs = _ti_ioc_queryset(request.user, request.GET)

    rows = (
        qs.extra(select={'day': "DATE(created_at)"})
          .values('day')
          .annotate(count=Count('id'))
          .order_by('day')
    )

    days = {str(r['day']): r['count'] for r in rows}

    # Fill all days in range so JS gets a complete grid
    date_from_str = request.GET.get('date_from', '')
    date_to_str   = request.GET.get('date_to',   '')
    try:
        d_from = date.fromisoformat(date_from_str)
    except ValueError:
        d_from = date.today() - timedelta(days=364)
    try:
        d_to = date.fromisoformat(date_to_str)
    except ValueError:
        d_to = date.today()

    result = []
    cur = d_from
    while cur <= d_to:
        key = cur.isoformat()
        result.append({'date': key, 'count': days.get(key, 0)})
        cur += timedelta(days=1)

    return JsonResponse({'days': result})


@login_required(login_url='login')
def api_ti_pivot(request):
    """
    GET /api/threat-intel/pivot/?value=<ioc_value>
    Returns all accessible incidents that contain an IOC with this exact value,
    plus aggregated reputation data.
    """
    value = request.GET.get('value', '').strip()
    if not value:
        return JsonResponse({'error': 'value required'}, status=400)

    incidents = _accessible_incidents(request.user)
    iocs = (
        GenericIoc.objects
        .filter(incident__in=incidents, value=value)
        .select_related('incident')
        .order_by('-created_at')
    )

    results = []
    best_rep = None
    SEVERITY = {'malicious': 3, 'suspicious': 2, 'clean': 1, 'unknown': 0}
    for ioc in iocs:
        rep = ioc.reputation or {}
        verdict = rep.get('status', 'unknown')
        if best_rep is None or SEVERITY.get(verdict, 0) > SEVERITY.get(best_rep.get('status', 'unknown'), 0):
            best_rep = rep
        results.append({
            'ioc_id':      ioc.id,
            'incident_id': ioc.incident_id,
            'incident':    ioc.incident.name,
            'severity':    ioc.incident.severity,
            'type':        ioc.type,
            'status':      ioc.status,
            'verdict':     verdict,
            'created_at':  ioc.created_at.isoformat(),
        })

    return JsonResponse({
        'value':      value,
        'count':      len(results),
        'reputation': best_rep,
        'incidents':  results,
    })


# ── Campaign CRUD ─────────────────────────────────────────────────────────────

@login_required(login_url='login')
def api_campaigns_list(request):
    """GET /api/campaigns/ — list campaigns with summary stats."""
    from django.db.models import Count

    accessible_ids = list(_accessible_incidents(request.user).values_list('id', flat=True))
    campaigns = Campaign.objects.prefetch_related('incidents').order_by('-created_at')

    data = []
    for c in campaigns:
        accessible_incidents = c.incidents.filter(id__in=accessible_ids).select_related()
        incident_ids = [i.id for i in accessible_incidents]
        ioc_count = GenericIoc.objects.filter(incident_id__in=incident_ids).count()
        data.append({
            'id':           c.id,
            'name':         c.name,
            'description':  c.description,
            'color':        c.color,
            'start_date':   c.start_date.isoformat() if c.start_date else None,
            'end_date':     c.end_date.isoformat() if c.end_date else None,
            'incident_count': len(incident_ids),
            'ioc_count':    ioc_count,
            'created_at':   c.created_at.isoformat(),
            'incidents':    [{'id': i.id, 'name': i.name, 'severity': i.severity, 'status': i.status}
                             for i in accessible_incidents],
        })

    return JsonResponse({'campaigns': data})


_HEX_COLOR_RE = re.compile(r'^#[0-9a-fA-F]{6}$')

def _safe_color(val, default='#c49840'):
    """Validate a CSS hex color string before storing."""
    if val and _HEX_COLOR_RE.match(str(val)):
        return str(val)
    return default

def _safe_date(val):
    """Parse a YYYY-MM-DD string or return None; never raises."""
    if not val:
        return None
    try:
        from datetime import date as _date
        return _date.fromisoformat(str(val)[:10])
    except ValueError:
        return None

def _can_mutate_campaign(user, campaign):
    """True if user may edit/delete this campaign (creator or superuser)."""
    return user.is_superuser or campaign.created_by_id == user.pk


@login_required(login_url='login')
def api_campaign_create(request):
    """POST /api/campaigns/create/"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed.'}, status=405)
    try:
        body = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON.'}, status=400)

    name = body.get('name', '').strip()
    if not name:
        return JsonResponse({'error': 'name is required.'}, status=400)

    campaign = Campaign.objects.create(
        name=name[:100],
        description=body.get('description', '')[:2000],
        color=_safe_color(body.get('color')),
        created_by=request.user,
        start_date=_safe_date(body.get('start_date')),
        end_date=_safe_date(body.get('end_date')),
    )

    incident_ids = body.get('incident_ids', [])
    if incident_ids:
        accessible = _accessible_incidents(request.user).filter(id__in=incident_ids)
        campaign.incidents.set(accessible)

    return JsonResponse({'status': 'created', 'id': campaign.id, 'name': campaign.name})


@login_required(login_url='login')
def api_campaign_update(request, campaign_id):
    """POST /api/campaigns/<id>/update/"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed.'}, status=405)
    campaign = get_object_or_404(Campaign, pk=campaign_id)
    if not _can_mutate_campaign(request.user, campaign):
        return JsonResponse({'error': 'Only the campaign creator can update it.'}, status=403)
    try:
        body = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON.'}, status=400)

    if 'name' in body:
        campaign.name = (body['name'].strip() or campaign.name)[:100]
    if 'description' in body:
        campaign.description = body['description'][:2000]
    if 'color' in body:
        campaign.color = _safe_color(body['color'], campaign.color)
    if 'start_date' in body:
        campaign.start_date = _safe_date(body['start_date'])
    if 'end_date' in body:
        campaign.end_date = _safe_date(body['end_date'])
    campaign.save()

    if 'incident_ids' in body:
        accessible = _accessible_incidents(request.user).filter(id__in=body['incident_ids'])
        campaign.incidents.set(accessible)

    return JsonResponse({'status': 'updated', 'id': campaign.id})


@login_required(login_url='login')
def api_campaign_delete(request, campaign_id):
    """POST /api/campaigns/<id>/delete/"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed.'}, status=405)
    campaign = get_object_or_404(Campaign, pk=campaign_id)
    if not _can_mutate_campaign(request.user, campaign):
        return JsonResponse({'error': 'Only the campaign creator can delete it.'}, status=403)
    campaign.delete()
    return JsonResponse({'status': 'deleted'})


@login_required(login_url='login')
def api_campaign_add_incident(request, campaign_id):
    """POST /api/campaigns/<id>/add-incident/  Body: { incident_id }"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed.'}, status=405)
    campaign = get_object_or_404(Campaign, pk=campaign_id)
    try:
        body = json.loads(request.body)
        inc_id = int(body['incident_id'])
    except (json.JSONDecodeError, KeyError, ValueError, TypeError):
        return JsonResponse({'error': 'incident_id required.'}, status=400)

    accessible = _accessible_incidents(request.user)
    if not accessible.filter(pk=inc_id).exists():
        return JsonResponse({'error': 'Incident not found or not accessible.'}, status=404)

    campaign.incidents.add(inc_id)
    return JsonResponse({'status': 'added'})


@login_required(login_url='login')
def api_campaign_remove_incident(request, campaign_id):
    """POST /api/campaigns/<id>/remove-incident/  Body: { incident_id }"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed.'}, status=405)
    campaign = get_object_or_404(Campaign, pk=campaign_id)
    try:
        body = json.loads(request.body)
        inc_id = int(body['incident_id'])
    except (json.JSONDecodeError, KeyError, ValueError, TypeError):
        return JsonResponse({'error': 'incident_id required.'}, status=400)

    # Verify the user can see the incident before removing it from the campaign
    if not _accessible_incidents(request.user).filter(pk=inc_id).exists():
        return JsonResponse({'error': 'Incident not found or not accessible.'}, status=404)

    campaign.incidents.remove(inc_id)
    return JsonResponse({'status': 'removed'})


# ── Exports ───────────────────────────────────────────────────────────────────

@login_required(login_url='login')
def api_ti_export_csv(request):
    """GET /api/threat-intel/export/csv/ — export filtered IOCs as CSV."""
    import csv as csv_mod

    qs = _ti_ioc_queryset(request.user, request.GET).select_related('incident').order_by('-created_at')

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="threat_intel_iocs.csv"'

    writer = csv_mod.writer(response)
    writer.writerow(['value', 'type', 'status', 'verdict', 'vt_malicious', 'vt_suspicious',
                     'vt_total', 'incident', 'incident_id', 'created_at', 'description'])
    for ioc in qs:
        rep = ioc.reputation or {}
        vt  = rep.get('vt', {}) or {}
        writer.writerow([
            ioc.value,
            ioc.type,
            ioc.status,
            rep.get('status', ''),
            vt.get('malicious', ''),
            vt.get('suspicious', ''),
            vt.get('total', ''),
            ioc.incident.name,
            ioc.incident_id,
            ioc.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            ioc.description,
        ])
    return response


@login_required(login_url='login')
def api_ti_export_json(request):
    """GET /api/threat-intel/export/json/ — export filtered IOCs as JSON."""
    qs = _ti_ioc_queryset(request.user, request.GET).select_related('incident').order_by('-created_at')

    data = []
    for ioc in qs:
        rep = ioc.reputation or {}
        data.append({
            'value':       ioc.value,
            'type':        ioc.type,
            'status':      ioc.status,
            'reputation':  rep,
            'incident':    {'id': ioc.incident_id, 'name': ioc.incident.name},
            'created_at':  ioc.created_at.isoformat(),
            'description': ioc.description,
        })

    content = json.dumps({'iocs': data, 'total': len(data)}, indent=2)
    response = HttpResponse(content, content_type='application/json')
    response['Content-Disposition'] = 'attachment; filename="threat_intel_iocs.json"'
    return response


@login_required(login_url='login')
def api_ti_export_pdf(request):
    """GET /api/threat-intel/export/pdf/ — export intelligence report as PDF."""
    from django.utils import timezone as tz

    qs  = _ti_ioc_queryset(request.user, request.GET).select_related('incident').order_by('-created_at')
    iocs = list(qs[:500])

    total   = len(iocs)
    mal_cnt = sum(1 for i in iocs if (i.reputation or {}).get('status') == 'malicious')
    susp_cnt = sum(1 for i in iocs if (i.reputation or {}).get('status') == 'suspicious')

    # Counts by type
    type_counts: dict = {}
    for ioc in iocs:
        type_counts[ioc.type] = type_counts.get(ioc.type, 0) + 1

    # Counts by incident
    inc_counts: dict = {}
    for ioc in iocs:
        key = ioc.incident.name
        inc_counts[key] = inc_counts.get(key, 0) + 1

    # Active filters for context
    filters = {k: v for k, v in request.GET.items() if v}

    html_string = render_to_string('threat_intel_report.html', {
        'iocs':        iocs,
        'total':       total,
        'malicious':   mal_cnt,
        'suspicious':  susp_cnt,
        'type_counts': sorted(type_counts.items(), key=lambda x: -x[1]),
        'inc_counts':  sorted(inc_counts.items(), key=lambda x: -x[1]),
        'filters':     filters,
        'generated_at': tz.now(),
        'user':        request.user,
    })

    buffer = BytesIO()
    pisa.CreatePDF(html_string, dest=buffer)
    buffer.seek(0)

    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="threat_intel_report.pdf"'
    return response


@login_required(login_url='login')
def api_ti_campaign_stats(request):
    """GET /api/threat-intel/campaign-stats/ — incident stats for the campaign tab."""
    from django.db.models import Count
    from django.db.models.functions import TruncMonth

    incidents = _accessible_incidents(request.user)

    date_from = request.GET.get('date_from', '').strip()
    date_to   = request.GET.get('date_to', '').strip()

    if date_from:
        try:
            incidents = incidents.filter(created_at__date__gte=date_from)
        except Exception:
            pass
    if date_to:
        try:
            incidents = incidents.filter(created_at__date__lte=date_to)
        except Exception:
            pass

    from django.db.models import Avg

    total     = incidents.count()
    open_     = incidents.filter(status='OPEN').count()
    in_prog   = incidents.filter(status='IN_PROGRESS').count()
    resolved  = incidents.filter(status='RESOLVED').count()
    closed    = incidents.filter(status='CLOSED').count()
    crit_high = incidents.filter(severity__in=['CRITICAL', 'HIGH']).count()
    camp_cnt  = Campaign.objects.count()

    # SLA metrics — only from incidents that have been resolved/closed
    closed_qs = incidents.filter(status__in=['RESOLVED', 'CLOSED'])
    sla = closed_qs.aggregate(
        avg_ttd=Avg('time_to_detect'),
        avg_ttr=Avg('time_to_respond'),
        avg_dur=Avg('duration'),
    )

    def fmt_td(td):
        if not td:
            return None
        secs = int(td.total_seconds())
        if secs <= 0:
            return None
        h, m = divmod(secs // 60, 60)
        d, h = divmod(h, 24)
        if d:
            return f"{d}d {h}h"
        return f"{h}h {m:02d}m"

    monthly_qs = (
        incidents
        .annotate(month=TruncMonth('created_at'))
        .values('month')
        .annotate(count=Count('id'))
        .order_by('month')
    )
    monthly = [
        {'month': row['month'].strftime('%Y-%m'), 'count': row['count']}
        for row in monthly_qs
        if row['month']
    ]

    # Resolution breakdown for closed/resolved incidents
    from django.db.models import Count as _Count
    resolution_qs = (
        incidents.exclude(resolution='')
        .values('resolution')
        .annotate(count=_Count('id'))
        .order_by('-count')
    )
    resolution_breakdown = [
        {'resolution': r['resolution'], 'count': r['count']}
        for r in resolution_qs
    ]
    tp_count = next((r['count'] for r in resolution_breakdown if r['resolution'] == 'TRUE_POSITIVE'), 0)
    closed_total = resolved + closed
    tp_rate = round(tp_count / closed_total * 100) if closed_total else None

    return JsonResponse({
        'total':       total,
        'open':        open_,
        'in_progress': in_prog,
        'resolved':    resolved,
        'closed':      closed,
        'crit_high':   crit_high,
        'campaigns':   camp_cnt,
        'monthly':     monthly,
        'avg_ttd':     fmt_td(sla['avg_ttd']),
        'avg_ttr':     fmt_td(sla['avg_ttr']),
        'avg_duration': fmt_td(sla['avg_dur']),
        'resolution_breakdown': resolution_breakdown,
        'tp_rate':     tp_rate,
    })


@login_required(login_url='login')
def api_ti_campaign_report_pdf(request):
    """GET /api/threat-intel/export/campaign-pdf/ — management incident report PDF."""
    from django.utils import timezone as tz
    from django.db.models import Count
    from django.db.models.functions import TruncMonth

    incidents = _accessible_incidents(request.user)

    date_from = request.GET.get('date_from', '').strip()
    date_to   = request.GET.get('date_to', '').strip()

    if date_from:
        try:
            incidents = incidents.filter(created_at__date__gte=date_from)
        except Exception:
            pass
    if date_to:
        try:
            incidents = incidents.filter(created_at__date__lte=date_to)
        except Exception:
            pass

    if date_from and date_to:
        period_label = f"{date_from} → {date_to}"
    elif date_from:
        period_label = f"From {date_from}"
    elif date_to:
        period_label = f"Until {date_to}"
    else:
        period_label = 'All time'

    incidents_list = list(
        incidents
        .prefetch_related('campaigns')
        .order_by('-created_at')[:200]
    )

    total     = len(incidents_list)
    open_cnt  = sum(1 for i in incidents_list if i.status == 'OPEN')
    inprog    = sum(1 for i in incidents_list if i.status == 'IN_PROGRESS')
    resolved  = sum(1 for i in incidents_list if i.status in ('RESOLVED', 'CLOSED'))
    crit_high = sum(1 for i in incidents_list if i.severity in ('CRITICAL', 'HIGH'))

    campaigns = Campaign.objects.prefetch_related('incidents').order_by('name')

    monthly_qs = (
        incidents
        .annotate(month=TruncMonth('created_at'))
        .values('month')
        .annotate(count=Count('id'))
        .order_by('month')
    )
    monthly = [
        {'month': row['month'].strftime('%b %Y'), 'count': row['count']}
        for row in monthly_qs
        if row['month']
    ]
    max_monthly = max((m['count'] for m in monthly), default=1)

    # SLA from resolved/closed incidents in the period
    from django.db.models import Avg
    closed_qs = incidents.filter(status__in=['RESOLVED', 'CLOSED'])
    sla = closed_qs.aggregate(
        avg_ttd=Avg('time_to_detect'),
        avg_ttr=Avg('time_to_respond'),
        avg_dur=Avg('duration'),
    )

    def fmt_td_pdf(td):
        if not td:
            return '—'
        secs = int(td.total_seconds())
        if secs <= 0:
            return '—'
        h, m = divmod(secs // 60, 60)
        d, h = divmod(h, 24)
        if d:
            return f"{d}d {h}h"
        return f"{h}h {m:02d}m"

    # Resolution breakdown
    res_qs = (
        incidents.exclude(resolution='')
        .values('resolution')
        .annotate(count=Count('id'))
        .order_by('-count')
    )
    resolution_breakdown = [{'resolution': r['resolution'], 'count': r['count']} for r in res_qs]
    closed_total = sum(r['count'] for r in resolution_breakdown)
    tp_count = next((r['count'] for r in resolution_breakdown if r['resolution'] == 'TRUE_POSITIVE'), 0)
    tp_rate = round(tp_count / closed_total * 100) if closed_total else None
    res_max = max((r['count'] for r in resolution_breakdown), default=1)

    RESOLUTION_LABELS = dict(Incident.RESOLUTION_CHOICES)

    html = render_to_string('threat_intel_campaign_report.html', {
        'incidents':            incidents_list,
        'campaigns':            campaigns,
        'total':                total,
        'open':                 open_cnt,
        'in_progress':          inprog,
        'resolved':             resolved,
        'crit_high':            crit_high,
        'period_label':         period_label,
        'date_from':            date_from,
        'date_to':              date_to,
        'generated_at':         tz.now(),
        'user':                 request.user,
        'monthly':              monthly,
        'max_monthly':          max_monthly,
        'avg_ttd':              fmt_td_pdf(sla['avg_ttd']),
        'avg_ttr':              fmt_td_pdf(sla['avg_ttr']),
        'avg_duration':         fmt_td_pdf(sla['avg_dur']),
        'resolution_breakdown': resolution_breakdown,
        'res_max':              res_max,
        'tp_rate':              tp_rate,
        'RESOLUTION_LABELS':    RESOLUTION_LABELS,
    })

    buffer = BytesIO()
    pisa.CreatePDF(html, dest=buffer)
    buffer.seek(0)

    # Sanitise date strings before embedding in HTTP header
    safe_from = re.sub(r'[^0-9\-]', '', date_from or 'all')[:10]
    safe_to   = re.sub(r'[^0-9\-]', '', date_to   or 'all')[:10]
    fname = f"incident_report_{safe_from}_{safe_to}.pdf"
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{fname}"'
    return response


# ─────────────────────────────────────────────────────────────────────────────
# BATCH INCIDENT OPERATIONS
# ─────────────────────────────────────────────────────────────────────────────

_VALID_STATUSES    = {'OPEN', 'IN_PROGRESS', 'RESOLVED', 'CLOSED'}
_VALID_SEVERITIES  = {'LOW', 'MEDIUM', 'HIGH', 'CRITICAL'}
_VALID_RESOLUTIONS = {r[0] for r in Incident.RESOLUTION_CHOICES}


@login_required(login_url='login')
def api_incidents_batch(request):
    """
    POST /api/incidents/batch/
    Body: { ids: [1,2,3], action: 'status'|'severity'|'resolution'|
                                   'assign_campaign'|'remove_campaign'|
                                   'assign_categories'|'merge'|'delete',
            ...action-specific fields... }
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required.'}, status=405)
    try:
        body = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON.'}, status=400)

    raw_ids = body.get('ids', [])
    try:
        ids = [int(i) for i in raw_ids]
    except (ValueError, TypeError):
        return JsonResponse({'error': 'ids must be integers.'}, status=400)

    if not ids:
        return JsonResponse({'error': 'No incidents selected.'}, status=400)
    if len(ids) > 200:
        return JsonResponse({'error': 'Batch limited to 200 incidents.'}, status=400)

    action = body.get('action', '')

    # Scope to incidents the user can access
    accessible = _accessible_incidents(request.user)
    incidents  = accessible.filter(id__in=ids)

    # For mutations that need INCIDENT_LEAD on each incident, we check per-incident
    def _is_lead(inc):
        return (request.user.is_superuser or
                UserRole.objects.filter(user=request.user, incident=inc, role='INCIDENT_LEAD').exists())

    # ── Status change ─────────────────────────────────────────────────────────
    if action == 'status':
        new_status = body.get('status', '').strip().upper()
        if new_status not in _VALID_STATUSES:
            return JsonResponse({'error': 'Invalid status.'}, status=400)

        if new_status == 'CLOSED':
            resolution = body.get('resolution', '').strip()
            if not resolution:
                return JsonResponse({'error': 'resolution_required', 'message': 'Resolution required to close.'}, status=400)
            if resolution not in _VALID_RESOLUTIONS:
                return JsonResponse({'error': 'Invalid resolution.'}, status=400)
            resolution_note = body.get('resolution_note', '')[:2000]
            incidents.update(status=new_status, resolution=resolution, resolution_note=resolution_note)
        else:
            incidents.update(status=new_status)

        return JsonResponse({'status': 'ok', 'updated': incidents.count()})

    # ── Severity change ───────────────────────────────────────────────────────
    if action == 'severity':
        new_sev = body.get('severity', '').strip().upper()
        if new_sev not in _VALID_SEVERITIES:
            return JsonResponse({'error': 'Invalid severity.'}, status=400)
        incidents.update(severity=new_sev)
        return JsonResponse({'status': 'ok', 'updated': incidents.count()})

    # ── Resolution update (without status change) ─────────────────────────────
    if action == 'resolution':
        resolution = body.get('resolution', '').strip()
        if not resolution:
            incidents.update(resolution='', resolution_note='')
        elif resolution in _VALID_RESOLUTIONS:
            incidents.update(resolution=resolution, resolution_note=body.get('resolution_note', '')[:2000])
        else:
            return JsonResponse({'error': 'Invalid resolution.'}, status=400)
        return JsonResponse({'status': 'ok', 'updated': incidents.count()})

    # ── Assign to campaign ────────────────────────────────────────────────────
    if action == 'assign_campaign':
        try:
            cid = int(body['campaign_id'])
        except (KeyError, ValueError, TypeError):
            return JsonResponse({'error': 'campaign_id required.'}, status=400)
        campaign = get_object_or_404(Campaign, pk=cid)
        for inc in incidents:
            campaign.incidents.add(inc)
        return JsonResponse({'status': 'ok', 'updated': incidents.count()})

    # ── Remove from campaign ──────────────────────────────────────────────────
    if action == 'remove_campaign':
        try:
            cid = int(body['campaign_id'])
        except (KeyError, ValueError, TypeError):
            return JsonResponse({'error': 'campaign_id required.'}, status=400)
        campaign = get_object_or_404(Campaign, pk=cid)
        for inc in incidents:
            campaign.incidents.remove(inc)
        return JsonResponse({'status': 'ok', 'updated': incidents.count()})

    # ── Assign categories ──────────────────────────────────────────────────────
    if action == 'assign_categories':
        cat_ids = body.get('category_ids', [])
        try:
            cat_ids = [int(c) for c in cat_ids]
        except (ValueError, TypeError):
            return JsonResponse({'error': 'category_ids must be integers.'}, status=400)
        cats = IncidentCategory.objects.filter(id__in=cat_ids)
        for inc in incidents:
            for cat in cats:
                inc.categories.add(cat)
        return JsonResponse({'status': 'ok', 'updated': incidents.count()})

    # ── Merge all into primary ────────────────────────────────────────────────
    if action == 'merge':
        try:
            primary_id = int(body['primary_id'])
        except (KeyError, ValueError, TypeError):
            return JsonResponse({'error': 'primary_id required.'}, status=400)
        if primary_id not in ids:
            return JsonResponse({'error': 'primary_id must be one of the selected incidents.'}, status=400)

        primary = get_object_or_404(Incident, pk=primary_id)
        if not _is_lead(primary):
            return JsonResponse({'error': 'INCIDENT_LEAD role required on the primary incident.'}, status=403)

        targets = incidents.exclude(pk=primary_id)
        merged = 0
        for target in targets:
            target.genericiocs.all().update(incident=primary)
            target.actions.all().update(incident=primary)
            target.notes.all().update(incident=primary)
            target.tasks.all().update(incident=primary)
            target.impacts.all().update(incident=primary)
            target.messages.all().update(incident=primary)
            target.shared_files.all().update(incident=primary)
            target.audit_logs.all().update(incident=primary)
            target.tags.all().update(incident=primary)
            for cat in target.categories.all():
                primary.categories.add(cat)
            for camp in target.campaigns.all():
                camp.incidents.add(primary)
            for ur in target.incident_roles.all():
                if not UserRole.objects.filter(user=ur.user, incident=primary).exists():
                    UserRole.objects.create(user=ur.user, incident=primary, role=ur.role)
            target.delete()
            merged += 1

        return JsonResponse({'status': 'ok', 'merged': merged, 'primary_id': primary_id})

    # ── Delete ────────────────────────────────────────────────────────────────
    if action == 'delete':
        deleted = 0
        for inc in list(incidents):
            if _is_lead(inc):
                inc.delete()
                deleted += 1
        if deleted == 0:
            return JsonResponse({'error': 'No incidents deleted. INCIDENT_LEAD role required.'}, status=403)
        return JsonResponse({'status': 'ok', 'deleted': deleted})

    return JsonResponse({'error': f'Unknown action: {action}'}, status=400)


# ─────────────────────────────────────────────────────────────────────────────
# MERGE INCIDENTS
# ─────────────────────────────────────────────────────────────────────────────

@login_required(login_url='login')
@verify_permissions(['INCIDENT_LEAD'])
def merge_incident(request, id):
    """
    POST /api/incident/<id>/merge/
    Body: { target_id: <int> }

    Moves all data from <target_id> into <id> (primary), then deletes the target.
    Requires INCIDENT_LEAD on the primary incident; user must also be able to see
    the target incident (any role or public).
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required.'}, status=405)

    try:
        body       = json.loads(request.body)
        target_id  = int(body['target_id'])
    except (json.JSONDecodeError, KeyError, ValueError, TypeError):
        return JsonResponse({'error': 'target_id required.'}, status=400)

    if target_id == id:
        return JsonResponse({'error': 'Cannot merge an incident with itself.'}, status=400)

    primary = get_object_or_404(Incident, pk=id)
    target  = get_object_or_404(Incident, pk=target_id)

    # User must be able to see the target incident
    accessible = _accessible_incidents(request.user)
    if not accessible.filter(pk=target_id).exists():
        return JsonResponse({'error': 'Target incident not found or not accessible.'}, status=404)

    # ── Reassign all related objects ──────────────────────────────────────────

    # Simple FK reassignments
    target.genericiocs.all().update(incident=primary)
    target.actions.all().update(incident=primary)
    target.notes.all().update(incident=primary)
    target.tasks.all().update(incident=primary)
    target.impacts.all().update(incident=primary)
    target.messages.all().update(incident=primary)
    target.shared_files.all().update(incident=primary)
    target.audit_logs.all().update(incident=primary)

    # Tags — belong to the incident, move them over
    target.tags.all().update(incident=primary)

    # Categories — merge M2M
    for cat in target.categories.all():
        primary.categories.add(cat)

    # UserRoles — add target's members to primary if not already there
    ROLE_RANK = {'INCIDENT_LEAD': 3, 'RESPONDER': 2, 'READER': 1, 'PUBLIC_VIEWER': 0}
    existing_roles = {ur.user_id: ur for ur in primary.incident_roles.all()}
    for ur in target.incident_roles.all():
        if ur.user_id not in existing_roles:
            ur.incident = primary
            ur.save()
        else:
            # Keep the higher-privilege role
            if ROLE_RANK.get(ur.role, 0) > ROLE_RANK.get(existing_roles[ur.user_id].role, 0):
                existing = existing_roles[ur.user_id]
                existing.role = ur.role
                existing.save()

    # Campaigns — re-point to primary
    for campaign in target.campaigns.all():
        campaign.incidents.add(primary)
        campaign.incidents.remove(target)

    # ── Audit log on primary ──────────────────────────────────────────────────
    _audit(primary, request, 'UPDATE', 'Incident',
           f'Merged incident #{target.id} "{target.name}" into this incident. '
           f'All IOCs, actions, notes, tasks, impacts and roles transferred.')

    target_name = target.name
    target.delete()

    return JsonResponse({
        'status':      'merged',
        'primary_id':  primary.id,
        'target_name': target_name,
        'message':     f'"{target_name}" has been merged into "{primary.name}".',
    })
